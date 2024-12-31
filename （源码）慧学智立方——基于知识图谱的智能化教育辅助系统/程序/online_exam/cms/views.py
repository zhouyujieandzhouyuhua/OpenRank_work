import ast
from .models import Score
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from itertools import zip_longest
from django.core.paginator import Paginator
from django.http import JsonResponse, HttpResponseRedirect
from django.shortcuts import render
from .models import Record, TestPaper, Question,StudentPaper, QuizQuestion, KnowledgeNode
from user.models import Subject
import random
from util.mypage import Pagination
from django.core.files.storage import default_storage
from flask import Flask, request, render_template
import os
import requests
import docx
from sparkai.llm.llm import ChatSparkLLM, ChunkPrintHandler
from sparkai.core.messages import ChatMessage
from pathlib import Path
from django.core.files.storage import default_storage
from django.http import JsonResponse
from django.conf import settings
from builtins import Exception, str, bytes
import xml.etree.ElementTree as ET
import websocket
import datetime
import hashlib
import base64
import hmac
import json
from urllib.parse import urlencode
import time
import ssl
from wsgiref.handlers import format_date_time
from datetime import datetime
from time import mktime
import _thread as thread
from django.views import View
import uuid
import subprocess
from django.db.models import Count

from django.shortcuts import render,redirect
from .models import Course,News,IndividualVideo
import openpyxl
from django.shortcuts import get_object_or_404
from django.utils import timezone
from django.core.exceptions import ObjectDoesNotExist
import random
from django.urls import resolve
from hanlp_restful import HanLPClient
# from .models import Topic, Board,Comment
HanLP = HanLPClient('https://www.hanlp.com/api', auth='NTg4NEBiYnMuaGFubHAuY29tOlZMV0tXSGNMdzNobWRlb0Q=', language='zh') # auth不填则匿名，zh中文，mul多语种   #,verify=False
from .models import Topic,Comment,IndividualTimu

from django.views.decorators.csrf import csrf_exempt
from docx import Document
import requests
from bs4 import BeautifulSoup
import pandas as pd
import sqlite3
import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import mean_absolute_error, mean_squared_error
import os
from django.http import HttpResponse
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
from .models import IndividualTimu
from docx.shared import Pt, RGBColor
from itertools import islice
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import oss2
from django.http import JsonResponse
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
from django.core.mail import send_mail
import logging

from .models import Record

global global_predict_student_score

from .models import Curriculum, KnowledgePoint, Relationship



from py2neo import Graph, Node, Relationship


from .neo4j import Neo4jService



# SPARKAI_URL = 'wss://spark-api.xf-yun.com/v3.5/chat'
# SPARKAI_APP_ID = '4626a4ce'
# SPARKAI_API_SECRET = 'OTg0Y2E3MTIwOWU4MDc1YzQ1ZTE5MTgz'
# SPARKAI_API_KEY = 'f7a10584a7234e3de1cfcf49f76b1054'
# SPARKAI_DOMAIN = 'generalv3.5'


SPARKAI_URL = 'wss://spark-api.xf-yun.com/v3.5/chat'
SPARKAI_APP_ID = '218e4d0f'
SPARKAI_API_SECRET = 'NGE3ZTc4MzcxNDUxODE2ZjkyY2QzNzRm'
SPARKAI_API_KEY = '34e7ff29b289814db153ff4e11613dab'
SPARKAI_DOMAIN = 'generalv3.5'

spark = ChatSparkLLM(
    spark_api_url=SPARKAI_URL,
    spark_app_id=SPARKAI_APP_ID,
    spark_api_key=SPARKAI_API_KEY,
    spark_api_secret=SPARKAI_API_SECRET,
    spark_llm_domain=SPARKAI_DOMAIN,
    streaming=False,
)

# UPLOAD_FOLDER = 'D:\\Langchain-Chatchat\\yuwenzuowen'
save_dir = Path(settings.MEDIA_ROOT) / 'kouyugendu'
UPLOAD_FOLDER = Path(settings.MEDIA_ROOT) / 'yuwenzuowen'


def index(req):
    questions = len(Question.objects.all())
    shijuans = len(TestPaper.objects.all())
    role = req.session["role"]
    username = req.session['username']

    return render(req, 'index.html', locals())


def test_paper(request):
    shijuan_list = TestPaper.objects.all()
    role = request.session["role"]
    username = request.session['username']

    current_page = request.GET.get('page', 1)
    if not current_page:
        current_page = 1
    all_count = shijuan_list.count()
    page_obj = Pagination(current_page=current_page, all_count=all_count, per_page_num=6, pager_count=11)
    shijuan_querylist = shijuan_list[page_obj.start:page_obj.end]
    if request.is_ajax():
        delete_id = request.POST.get('delete_id')
        back_dict = {'code': 200, 'msg': ''}
        TestPaper.objects.filter(pk=delete_id).delete()
        back_dict['msg'] = '删除成功'
        return JsonResponse(back_dict)
    return render(request, 'test_paper.html', locals())


def login_out(req):
    try:
        del req.session['username']
        return HttpResponseRedirect('/')
    except Exception as e:
        return HttpResponseRedirect('/')


# Create your views here.
def login(req):
    return render(req, 'login.html')


def chengji(req):
    role = req.session["role"]
    username = req.session['username']


    return render(req, 'chengji.html', locals())


def get_chengji(request):
    """
    获取用户列表信息 | 模糊查询
    :param request:
    :return:
    """
    keyword = request.GET.get('name')
    page = request.GET.get("page", '')
    limit = request.GET.get("limit", '')
    response_data = {}
    response_data['code'] = 0
    response_data['msg'] = ''
    data = []
    if keyword is None:
        results = Record.objects.order_by("-grade").all()
        paginator = Paginator(results, limit)
        results = paginator.page(page)
        if results:
            for user in results:
                record = {
                    "id": user.id,
                    "name": user.name,
                    "number": user.xuehao,
                    "grade": user.grade,
                    'title': TestPaper.objects.filter(id=user.test_paper_id).first().title,
                    'create_time': user.exam_time.strftime('%Y-%m-%d %H:%m:%S'),
                }
                data.append(record)
            response_data['count'] = len(Record.objects.all())
            response_data['data'] = data
    else:
        users_all = Record.objects.filter(name__contains=keyword).all()
        paginator = Paginator(users_all, limit)
        results = paginator.page(page)
        if results:
            for user in results:
                record = {
                    "id": user.id,
                    "name": user.name,
                    "number": user.number,
                    "grade": user.grade,
                    'test_paper': TestPaper.objects.filter(id=user.test_paper_id).title,
                    'create_time': user.exam_time.strftime(
                        '%Y-%m-%d %H:%m:%S'),
                }
                data.append(record)
            response_data['count'] = len(users_all)
            response_data['data'] = data

    # print('response::::',response_data)

    return JsonResponse(response_data)


# ============================================================================================
def get_scores(request):
    scores = Score.objects.filter(name='周渝')  # 获取所有Score记录
    data = []
    for score in scores:
        row = {
            'student_number': score.student_number,
            'name': score.name,
            'subject': score.subject,
            'total_score': score.total_score,
            'time': score.time.strftime('%Y-%m-%d %H:%M:%S'),  # 格式化时间戳
            'question_1': score.question_1,
            'question_2': score.question_2,
            'question_3': score.question_3,
            'question_4': score.question_4,
            'question_5': score.question_5,
            'question_6': score.question_6,
            'question_7': score.question_7,
            'question_8': score.question_8,
            'question_9': score.question_9,
            'question_10': score.question_10,
            'question_11': score.question_11,
        }
        data.append(row)
        response_data = {
            "code": 0,  # 成功状态码
            "msg": "",  # 可选的消息
            "data": data  # 数据列表
        }

    return JsonResponse(response_data)


# =============================================================================================
def shijuan(req):
    username = req.session.get('username')
    role = req.session["role"]
    return render(req, 'shijuan.html', locals())


def get_shijuan(request):
    """
    获取用户列表信息 | 模糊查询
    :param request:
    :return:
    """
    keyword = request.GET.get('name')
    page = request.GET.get("page", '')
    limit = request.GET.get("limit", '')
    role_id = request.GET.get('position', '')
    response_data = {}
    response_data['code'] = 0
    response_data['msg'] = ''
    data = []
    if keyword is None:
        results = TestPaper.objects.all()
        paginator = Paginator(results, limit)
        results = paginator.page(page)
        if results:
            for user in results:
                record = {
                    "id": user.id,
                    "title": user.title,
                    "owner": user.owner,
                    "course": user.course,
                    'time': user.time,
                    'create_time': user.exam_time.strftime('%Y-%m-%d %H:%m:%S'),
                }
                data.append(record)
            response_data['count'] = len(TestPaper.objects.all())
            response_data['data'] = data
    else:
        users_all = TestPaper.objects.filter(name__contains=keyword).all()
        paginator = Paginator(users_all, limit)
        results = paginator.page(page)
        if results:
            for user in results:
                record = {
                    "id": user.id,
                    "title": user.title,
                    "owner": user.owner,
                    "course": user.course,
                    'time': user.time,
                    'create_time': user.exam_time.strftime('%Y-%m-%d %H:%m:%S'),
                }
                data.append(record)
            response_data['count'] = len(users_all)
            response_data['data'] = data
    return JsonResponse(response_data)


def shiti(req):
    username = req.session['username']
    role = req.session["role"]
    return render(req, 'question.html', locals())


def get_questions(request):
    """
    获取用户列表信息 | 模糊查询
    :param request:
    :return:
    """
    keyword = request.GET.get('title')
    page = request.GET.get("page", '')
    limit = request.GET.get("limit", '')
    response_data = {}
    response_data['code'] = 0
    response_data['msg'] = ''
    data = []
    if keyword is None:
        results = Question.objects.all()
        paginator = Paginator(results, limit)
        results = paginator.page(page)
        if results:
            for user in results:
                record = {
                    "id": user.id,
                    "title": user.title,
                    'course': Subject.objects.filter(id=user.course).first().subject_name,
                    "A": user.a,
                    "B": user.b,
                    "C": user.c,
                    'D': user.d,
                    'question_type': user.question_type,
                    'owner': user.owner,
                    'last_modify_time': user.last_modify_time.strftime('%Y-%m-%d %H:%m:%S'),
                    'answer': user.answer,
                    "difficulty": user.difficulty,
                    'score': user.score,
                }
                data.append(record)
            response_data['count'] = len(Question.objects.all())
            response_data['data'] = data
    else:
        users_all = Question.objects.filter(title__contains=keyword).all()
        paginator = Paginator(users_all, limit)
        results = paginator.page(page)
        if results:
            for user in results:
                record = {
                    "id": user.id,
                    "title": user.title,
                    'course': Subject.objects.filter(id=user.course).first().subject_name,
                    "A": user.a,
                    "B": user.b,
                    "C": user.c,
                    'D': user.d,
                    'question_type': user.question_type,
                    'owner': user.owner,
                    'last_modify_time': user.last_modify_time.strftime('%Y-%m-%d %H:%m:%S'),
                    'answer': user.answer,
                    "difficulty": user.difficulty,
                    'score': user.score,
                }
                data.append(record)
            response_data['count'] = len(users_all)
            response_data['data'] = data
    return JsonResponse(response_data)


def add_question(request):
    """
    添加员工
    :return:
    """
    try:
        title = request.POST.get('title')
        question_type = request.POST.get('question_type')
        score = request.POST.get('score')
        xxa, xxb, xxc, xxd = '', '', '', ''
        if question_type == '单选题':
            xxa = request.POST.get('xxa')
            xxb = request.POST.get('xxb')
            xxc = request.POST.get('xxc')
            xxd = request.POST.get('xxd')
        answer = request.POST.get('answer')

        difficulty = request.POST.get('difficulty')
        question = Question.objects.filter(title=title)
        if question:
            return JsonResponse({'message': '题目已存在,请直接登录'}, status=403)
        Question.objects.create(
            title=title,
            score=score,
            question_type=question_type,
            a=xxa,
            b=xxb,
            c=xxc,
            d=xxd,
            answer=answer,
            difficulty=difficulty,
            owner='admin',
            course=1
        )
        response_data = {'message': '注册成功'}
        return JsonResponse(response_data)
    except Exception as e:
        print(e)
        return JsonResponse({'message': '注册失败'}, status=401)


def inference(req):
    shijuan_id = req.POST.get('shijuan_id')
    ques1 = req.POST.get('question1')
    ques2 = req.POST.get('question2')
    ques3 = req.POST.get('question3')
    ques4 = req.POST.get('question4')
    ques5 = req.POST.get('question5')
    ques6 = req.POST.get('question6')
    ques7 = req.POST.get('question7')
    ques8 = req.POST.get('question8')
    ques9 = req.POST.get('question9')
    ques10 = req.POST.get('question10')
    ques11 = req.POST.get('question11')
    print(shijuan_id)
    answer = ast.literal_eval(TestPaper.objects.filter(id=shijuan_id).first().answer)
    score = 0
    if ques1 == answer[0]:
        score += 10
    if ques2 == answer[1]:
        score += 10
    if ques3 == answer[2]:
        score += 10
    if ques4 == answer[3]:
        score += 10
    if ques5 == answer[4]:
        score += 10
    if ques6 == answer[5]:
        score += 5
    if ques7 == answer[6]:
        score += 5
    if ques8 == answer[7]:
        score += 5
    if ques9 == answer[8]:
        score += 5
    if ques10 == answer[9]:
        score += 5
    if len(ques11) >= 30:
        score += 18
    elif len(ques11) > 15 and len(ques11) < 30:
        score += 10
    elif len(ques11) == 0:
        score += 0
    else:
        score += 5
    Record.objects.create(
        xuehao=1,
        name=req.session['username'],
        grade=score,
        test_paper_id=shijuan_id,
        answer=str([ques1, ques2, ques3, ques4, ques5, ques6, ques7, ques8, ques9, ques10, ques11]),
    )

    return JsonResponse({'msg': 'ok'})


# =====================================上传作业================================================================
def uphomework(request):
    if request.method == 'POST':
        file = request.FILES.get('file')
        if file:

            # 保存文件到指定路径
            file_path = default_storage.save('homework_photos/' + file.name, file)
            URL = "http://webapi.xfyun.cn/v1/service/v1/ocr/handwriting"

            APPID = "4626a4ce"

            API_KEY = "bc3519fad768aab956e9da66d95b2e9c"

            def getHeader():
                curTime = str(int(time.time()))
                param = "{\"language\":\"" + language + "\",\"location\":\"" + location + "\"}"
                paramBase64 = base64.b64encode(param.encode('utf-8'))

                m2 = hashlib.md5()
                str1 = API_KEY + curTime + str(paramBase64, 'utf-8')
                m2.update(str1.encode('utf-8'))
                checkSum = m2.hexdigest()
                # 组装http请求头
                header = {
                    'X-CurTime': curTime,
                    'X-Param': paramBase64,
                    'X-Appid': APPID,
                    'X-CheckSum': checkSum,
                    'Content-Type': 'application/x-www-form-urlencoded; charset=utf-8',
                }
                return header

            def getBody(filepath):
                with open(filepath, 'rb') as f:
                    imgfile = f.read()
                data = {'image': str(base64.b64encode(imgfile), 'utf-8')}
                return data

            # 语种设置
            language = "en"
            # 是否返回文本位置信息
            location = "true"
            # 图片上传接口地址
            # picFilePath = "D:/tools/ocr_handwriting_python3_demo/ocr_handwriting_python3.x_demo/2.jpg"
            picFilePath = file_path
            # headers=getHeader(language, location)
            r = requests.post(URL, headers=getHeader(), data=getBody(picFilePath))

            # json_str = b'{"code":"0","data":{"block":[{"type":"text","line":[{"confidence":1,"location":{"top_left":{"x":14,"y":63},"right_bottom":{"x":354,"y":159}},"word":[{"content":"Aa"},{"content":"Aa"}]},{"confidence":1,"location":{"top_left":{"x":319,"y":217},"right_bottom":{"x":353,"y":242}},"word":[{"content":"33"}]}]}]},"desc":"success","sid":"wcr00e44117@gz184019bdb4b8463000"}'
            json_str = r.content

            print('json_str', type(json_str))
            # 将字节字符串解码为普通字符串
            json_str = json_str.decode('utf-8')

            # 将字符串解析为字典
            data_dict = json.loads(json_str)

            # 定义一个函数来提取word字段的内容
            def extract_words(data):
                if isinstance(data, dict):
                    for key, value in data.items():
                        if key == 'word':
                            for word in value:
                                yield word['content']
                        else:
                            yield from extract_words(value)
                elif isinstance(data, list):
                    for item in data:
                        yield from extract_words(item)

            # 调用函数并打印结果
            words = list(extract_words(data_dict))
            out = ''  # out是原文
            for i in words:
                out = out + i
            print('out', out)
            return JsonResponse({"status": out})
        else:
            return JsonResponse({"status": "error"}, status=400)
    return JsonResponse({"status": "error"}, status=405)


def see_shijuan(req, pid):
    username = req.session.get('username')
    role = req.session["role"]
    shijuan = TestPaper.objects.filter(id=pid).first()
    xuanzeti = ast.literal_eval(shijuan.pid)[0:5]
    panduanti = ast.literal_eval(shijuan.pid)[5:10]
    jiandati = ast.literal_eval(shijuan.pid)[10:]
    xuanzeti_info = []
    panduanti_info = []
    jiandati_info = []
    shijuan_id = pid
    for i in xuanzeti:
        result = Question.objects.filter(id=i).first()
        xuanzeti_info.append(result)
    for j in panduanti:
        result = Question.objects.filter(id=j).first()
        panduanti_info.append(result)
    for k in jiandati:
        result = Question.objects.filter(id=k).first()
        jiandati_info.append(result)
    return render(req, 'see_shijuan.html', locals())


def to_test_paper(req, pid):
    username = req.session.get('username')

    role = req.session["role"]
    shijuan = TestPaper.objects.filter(id=pid).first()
    exam_time = shijuan.time * 60
    xuanzeti = ast.literal_eval(shijuan.pid)[0:7]
    panduanti = ast.literal_eval(shijuan.pid)[7:11]
    jiandati = ast.literal_eval(shijuan.pid)[11:12]
    xuanzeti_info = []
    panduanti_info = []
    jiandati_info = []
    shijuan_id = pid
    for i in xuanzeti:
        result = Question.objects.filter(id=i).first()
        xuanzeti_info.append(result)
    for j in panduanti:
        result = Question.objects.filter(id=j).first()
        panduanti_info.append(result)
    for k in jiandati:
        result = Question.objects.filter(id=k).first()
        jiandati_info.append(result)
    return render(req, 'timu.html', locals())


def edit_question(request):
    response_data = {}
    user_id = request.POST.get('id')
    title = request.POST.get('title')
    score = request.POST.get('score')
    question_type = request.POST.get('question_type')
    xxa, xxb, xxc, xxd = '', '', '', ''
    if question_type == '单选题':
        xxa = request.POST.get('xxa')
        xxb = request.POST.get('xxb')
        xxc = request.POST.get('xxc')
        xxd = request.POST.get('xxd')
    answer = request.POST.get('answer')
    difficulty = request.POST.get('difficulty')
    Question.objects.filter(id=user_id).update(
        title=title,
        score=score,
        question_type=question_type,
        a=xxa,
        b=xxb,
        c=xxc,
        d=xxd,
        answer=answer,
        difficulty=difficulty, )
    response_data['msg'] = 'success'
    return JsonResponse(response_data, status=201)


def edit_shijuan(request):
    response_data = {}
    user_id = request.POST.get('id')
    title = request.POST.get('name')
    exam_time = request.POST.get('exam_time')

    TestPaper.objects.filter(id=user_id).update(
        title=title,
        time=exam_time,
    )
    response_data['msg'] = 'success'
    return JsonResponse(response_data, status=201)


def del_question(request):
    question_id = request.POST.get('id')
    result = Question.objects.filter(id=question_id).first()
    try:
        if not result:
            response_data = {'error': '删除试题失败！', 'message': '找不到id为%s的试题' % question_id}
            return JsonResponse(response_data, status=403)
        result.delete()
        response_data = {'message': '删除成功！'}
        return JsonResponse(response_data, status=201)
    except Exception as e:
        response_data = {'message': '删除失败！'}
        return JsonResponse(response_data, status=403)


def add_shijuan(req):
    danxuantis = Question.objects.filter(question_type='单选题').all()
    panduantis = Question.objects.filter(question_type='判断题').all()
    jiandatis = Question.objects.filter(question_type='简答题').all()
    danxuanti_list = []
    answer = []
    panduanti_list = []
    jiandati_list = []
    try:
        for i in danxuantis:
            danxuanti_list.append(i.id)
        for j in panduantis:
            panduanti_list.append(j.id)
        for k in jiandatis:
            jiandati_list.append(k.id)
        jiandati = random.sample(jiandati_list, 1)
        danxuanti = random.sample(danxuanti_list, 5)
        panduanti = random.sample(panduanti_list, 5)
        shiti_list = danxuanti + panduanti + jiandati
        for i in shiti_list:
            shiti_answer = Question.objects.filter(id=i).first().answer
            answer.append(shiti_answer)

        TestPaper.objects.create(
            title=req.POST.get('name'),
            owner=req.POST.get('owner'),
            course='综合',
            time=req.POST.get('exam_time'),
            pid=str(danxuanti + panduanti + jiandati),
            answer=str(answer),
        )
        response_data = {'message': '新建成功'}
        return JsonResponse(response_data)
    except Exception as e:
        print(e)
        return JsonResponse({'message': '注册失败'}, status=401)


# =======================================作文批改========================================================
language = None


class ChangeLanguageView(View):  # 获取前端选择的作文批改语言
    def post(self, request):
        global language
        language = request.POST.get('language')
        print('当前处理作文的语言为:', language)
        return JsonResponse({'status': 'success'})


def upload_zuowen_photo(request):
    time.sleep(4)
    if request.method == 'POST':
        if 'photo' not in request.FILES:
            return JsonResponse({'error': '没有文件部分'}, status=400)
        file = request.FILES['photo']
        if not file.name:
            return JsonResponse({'error': '没有选择文件'}, status=400)

        # 确保上传目录存在
        if not os.path.exists(UPLOAD_FOLDER):
            os.makedirs(UPLOAD_FOLDER)

        file_path = os.path.join(UPLOAD_FOLDER, file.name)

        # 将文件保存到指定目录
        with open(file_path, 'wb+') as destination:
            for chunk in file.chunks():
                destination.write(chunk)



# -----------------------------------------暂时关闭了下面------------------------------------------------------------------------
        URL = "http://webapi.xfyun.cn/v1/service/v1/ocr/handwriting"
        APPID = "4626a4ce"
        API_KEY = "bc3519fad768aab956e9da66d95b2e9c"

        def getHeader():
            curTime = str(int(time.time()))
            param = "{\"language\":\"" + language + "\",\"location\":\"" + location + "\"}"
            paramBase64 = base64.b64encode(param.encode('utf-8'))

            m2 = hashlib.md5()
            str1 = API_KEY + curTime + str(paramBase64, 'utf-8')
            m2.update(str1.encode('utf-8'))
            checkSum = m2.hexdigest()
            # 组装http请求头
            header = {
                'X-CurTime': curTime,
                'X-Param': paramBase64,
                'X-Appid': APPID,
                'X-CheckSum': checkSum,
                'Content-Type': 'application/x-www-form-urlencoded; charset=utf-8',
            }
            return header

        def getBody(filepath):
            with open(filepath, 'rb') as f:
                imgfile = f.read()
            data = {'image': str(base64.b64encode(imgfile), 'utf-8')}
            return data

        # 是否返回文本位置信息
        location = "true"
        # 图片上传接口地址

        picFilePath = file_path
        print('路径：', picFilePath, '正在响应作文批改')
        print('当前处理的语言类型', language)
        # headers=getHeader(language, location)
        r = requests.post(URL, headers=getHeader(), data=getBody(picFilePath))

        json_str = r.content
        # 将字节字符串解码为普通字符串
        json_str = json_str.decode('utf-8')
        # 将字符串解析为字典
        data_dict = json.loads(json_str)

        # 定义一个函数来提取word字段的内容
        def extract_words(data):
            if isinstance(data, dict):
                for key, value in data.items():
                    if key == 'word':
                        for word in value:
                            yield word['content']
                    else:
                        yield from extract_words(value)
            elif isinstance(data, list):
                for item in data:
                    yield from extract_words(item)

        # 调用函数并打印结果
        words = list(extract_words(data_dict))
        out = ''  # out是原文
        for i in words:
            out = out + i
# ------------------------------------------暂时关闭了上面---------------------------------------------------------------------------------------

# ---------------------------------------------------暂时关闭了下面------------------------------------------------------------------------------------

        # ----------然后把out再调用文字批改api，得到out_xiugai--------------
        messages = [ChatMessage(
            role="user",
            content=out,
        )]
        handler = ChunkPrintHandler()
        a = spark.generate([messages], callbacks=[handler])
        first_reply_content = a.generations[0][0].message.content

# ---------------------------------------------------暂时关闭了上面------------------------------------------------------------------------------------





        return JsonResponse({'message': out, 'message_2': first_reply_content})
    else:
        return JsonResponse({'error': '无效的请求方法'}, status=405)


# =================================ai对话=================================================================
def handle_command_help():
    response = """
    同学你好呀！欢迎使用指令模式，你可以使用以下指令：
    1./command+智能组卷
    2./command+分析成绩
    """
    return response
def handle_command_shijuan(request):
    file_path = 'output_train_data.xlsx'
    data = pd.read_excel(file_path)

    # 数据预处理
    data = data.drop(columns=['Username'])
    X = data.drop(columns=['Average Exam Score'])
    y = data['Average Exam Score']

    # 分割数据
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

    # 标准化数据
    scaler = StandardScaler()
    X_train = scaler.fit_transform(X_train)
    X_test = scaler.transform(X_test)

    # 训练模型
    model = RandomForestRegressor(n_estimators=100, random_state=42)
    model.fit(X_train, y_train)

    # 评估模型
    y_pred = model.predict(X_test)
    mae = mean_absolute_error(y_test, y_pred)
    mse = mean_squared_error(y_test, y_pred)
    rmse = mse ** 0.5

    # print(f'Mean Absolute Error: {mae}')
    # print(f'Mean Squared Error: {mse}')
    # print(f'Root Mean Squared Error: {rmse}')

    # 预测新数据函数
    def predict_new_data(new_data):
        new_data_scaled = scaler.transform(new_data)
        predictions = model.predict(new_data_scaled)
        return predictions

    # 示例：预测最新的学生数据（即最后一行）
    last_student_data = X.tail(1)  # 选择DataFrame的最后一行
    last_student_predictions = predict_new_data(last_student_data)
    print(f'当前学生预测得分: {last_student_predictions[0]}')
    predict_student_score = int(last_student_predictions[0])

    if predict_student_score >= 90:
        file_url = 'https://rengongzhinengcourse.oss-cn-hangzhou.aliyuncs.com/personalized_exam%EF%BC%88A%E5%8D%B7%EF%BC%89.docx'
    elif predict_student_score >= 80:
        file_url = 'https://rengongzhinengcourse.oss-cn-hangzhou.aliyuncs.com/personalized_exam%20%28B%E5%8D%B7%29.docx'

    elif predict_student_score >= 70:
        file_url = 'https://rengongzhinengcourse.oss-cn-hangzhou.aliyuncs.com/personalized_exam%20%28%E9%9A%BE%E5%BA%A6C%29.docx'

    else:
        file_url = 'https://rengongzhinengcourse.oss-cn-hangzhou.aliyuncs.com/personalized_exam%20%28D%E5%8D%B7%29.docx'
    if request.method == 'POST':
        response_data = {
            'message': '咻~，属于你的个性化试卷生成完成啦！',
            'download_url': file_url,
        }
        return JsonResponse(response_data)
    else:
        return JsonResponse({'error': 'Invalid request method'}, status=400)



def handle_command_analysis():
    knowledge_points = {
        "1": "智能机器人系统绪论", "2": "智能机器人平台与系统", "3": "机器人运动", "4": "机器人控制",
        "5": "机器人传感器", "6": "机器人自主定位与建图", "7": "三维刚体运动的描述",
        "8": "贝叶斯滤波框架", "9": "同步定位与建图（SLAM）", "10": "视觉SLAM",
        "11": "多机器人编队控制", "12": "智能人机交互", "13": "人机共享控制",
        "14": "机器人操作系统—ROS重要工具", "15": "智能机器人关键技术", "16": "智能机器人典型案例"
    }

    file_path = 'output_train_data.xlsx'
    data = pd.read_excel(file_path)
    last_row_data = data.iloc[-1]

    # 分析播放进度和情感得分
    progress_columns = [f'Progress {i}' for i in range(1, 17)]
    emo_columns = [f'Emo {i}' for i in range(1, 17)]
    test_columns = [f'Test {i}' for i in range(1, 6)]

    under_30_progress = [i + 1 for i, val in enumerate(last_row_data[progress_columns]) if val < 30]
    under_60_emo = [i + 1 for i, val in enumerate(last_row_data[emo_columns]) if val < 60]
    test_scores = last_row_data[test_columns]
    avg_exam_score = last_row_data['Average_Exam_Score']

    # 创建Word文档
    doc = Document()
    doc.add_heading('成绩分析详细报告', 0)

    p = doc.add_paragraph()
    run = p.add_run('亲爱的同学你好！下面是你的成绩报告：')
    run.font.size = Pt(12)

    if under_30_progress:
        progress_str = '、'.join(map(str, under_30_progress))
        doc.add_paragraph(f'在该课程的播放进度中，你第 {progress_str} 节视频课不达标且播放进度小于30%。')
        doc.add_paragraph('这意味着你可能没有充分参与这些课程的学习。我们建议你回看这些视频，确保你理解了所有关键概念。')

    if under_60_emo:
        emo_str = '、'.join(map(str, under_60_emo))
        doc.add_paragraph(f'在这些视频课程中，你对于第 {emo_str} 节课程的情感分析评价低于60分。')
        points_str = '、'.join([knowledge_points[str(i)] for i in under_60_emo])
        doc.add_paragraph(f'这些课程分别对应的知识点是：{points_str}')
        doc.add_paragraph(
            '低情感得分可能表明你对这些主题的兴趣不高或感到困惑。尝试寻找额外资源，如在线教程或相关书籍，以增强对这些主题的理解。')

    test_scores_str = '、'.join(map(str, test_scores))
    doc.add_paragraph(f'在你的课后习题中，你的得分分别是：{test_scores_str}')
    if any(score < 60 for score in test_scores):
        doc.add_paragraph(
            '有些课后习题的得分较低，这可能意味着你在某些概念上需要更多练习。考虑安排复习时间和使用额外的习题集。')

    doc.add_paragraph(f'你近5次的考试成绩平均分是：{avg_exam_score}')
    if avg_exam_score < 60:
        doc.add_paragraph('你的平均考试成绩低于及格线。建议制定一个学习计划，重点复习薄弱环节，并寻求教师的帮助。')

    # 插入表格表示测试分数
    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '测试'
    for i in range(1, 6):
        hdr_cells[i].text = f'Test {i}'

    row_cells = table.add_row().cells
    row_cells[0].text = '分数'
    for i, score in enumerate(test_scores):
        row_cells[i + 1].text = str(score)

    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # 保存Word文档到本地
    local_file_path = r'C:\Users\1\Desktop\成绩分析报告.docx'
    doc.save(local_file_path)

    # 配置OSS
    accessKeyId = 'LTAI5tA5od2XVKZJcvWUhfsk'
    accessKeySecret = 'C768TfmVONkbvNdyzWvpHDFUlWB36l'
    auth = oss2.Auth(accessKeyId, accessKeySecret)
    endpoint = 'https://oss-cn-hangzhou.aliyuncs.com'
    bucketName = 'rengongzhinengcourse'
    bucket = oss2.Bucket(auth, endpoint, bucketName)

    # 上传文件到OSS
    objectName = '成绩分析报告.docx'
    bucket.put_object_from_file(objectName, local_file_path)

    # 生成下载链接
    fileLink = f'http://{bucketName}.oss-cn-hangzhou.aliyuncs.com/{objectName}'

    response_data = {
        'message': '咻~，属于你的成绩分析报告生成完成啦！',
        'download_url': fileLink,
    }

    return JsonResponse(response_data)

def ai_answer(request):
    if request.method == 'POST':
        # 从POST请求中获取问题,下面是拿的过程  先去html文件里命名然后再去填写url  再去更新数据库 和字段 字段要小心 记住这个流程！！
        question = request.POST.get('question')
        if question.startswith('/command'):
            # 提取命令
            command = question[len('/command'):].strip()

            # 在这里添加你要执行的命令处理逻辑
            if command == '+help':
                response = handle_command_help()

            elif command == '+智能组卷':
                return handle_command_shijuan(request)
            elif command == '+分析成绩':
                return handle_command_analysis()

            else:
                return JsonResponse({'response': '同学是不是忘了"+"号？输入指令要带上"+"号哦！请输入/command+help查看可用命令。'})

            return JsonResponse({'response': response})
        messages = [ChatMessage(
            role="user",
            content=question
        )]
        handler = ChunkPrintHandler()
        a = spark.generate([messages], callbacks=[handler])
        first_reply_content = a.generations[0][0].message.content
        return JsonResponse({'response': first_reply_content})
    else:
        return JsonResponse({'error': 'Invalid request method'}, status=400)


# ------------------------------------------------------------------------------------------------=
gendu_score = None


def gendu(request):
    if request.method == 'POST':
        audio_file = request.FILES['audio']
        sentence_index = request.POST['sentenceIndex']  # 此处我们获得了前端传来的句子de索引
        practice_level = request.POST['practiceLevel']  # 此处我们获得了前端传来的句子的练习等级
        sentence = request.POST['sentence']  # 此处我们获得了前端传来的句子
        save_dir = Path(settings.MEDIA_ROOT) / 'kouyugendu'
        if not save_dir.exists():
            save_dir.mkdir(parents=True, exist_ok=True)
        # 使用相对路径保存文件
        save_path = save_dir / audio_file.name
        with save_path.open('wb+') as destination:
            for chunk in audio_file.chunks():
                destination.write(chunk)
        # ===============================把文件拿出来调用a============================================================
        STATUS_FIRST_FRAME = 0  # 第一帧的标识
        STATUS_CONTINUE_FRAME = 1  # 中间帧标识
        STATUS_LAST_FRAME = 2  # 最后一帧的标识
        SUB = "ise"
        ENT = "en_vip"
        # 中文题型：read_syllable（单字朗读，汉语专有）read_word（词语朗读）read_sentence（句子朗读）read_chapter(篇章朗读)
        # 英文题型：read_word（词语朗读）read_sentence（句子朗读）read_chapter(篇章朗读)simple_expression（英文情景反应）read_choice（英文选择题）topic（英文自由题）retell（英文复述题）picture_talk（英文看图说话）oral_translation（英文口头翻译）
        CATEGORY = "read_sentence"
        TEXT = '\uFEFF' + sentence

        class Ws_Param(object):
            # 初始化
            def __init__(self, APPID, APIKey, APISecret, AudioFile, Text):
                self.APPID = APPID
                self.APIKey = APIKey
                self.APISecret = APISecret
                self.AudioFile = AudioFile
                self.Text = Text

                # 公共参数(common)
                self.CommonArgs = {"app_id": self.APPID}
                # 业务参数(business)，更多个性化参数可在官网查看
                self.BusinessArgs = {"category": CATEGORY, "sub": SUB, "ent": ENT, "cmd": "ssb",
                                     "auf": "audio/L16;rate=16000",
                                     "aue": "raw", "text": self.Text, "ttp_skip": True, "aus": 1}

            # 生成url
            def create_url(self):
                # wws请求对Python版本有要求，py3.7可以正常访问，如果py版本请求wss不通，可以换成ws请求，或者更换py版本
                url = 'ws://ise-api.xfyun.cn/v2/open-ise'
                # 生成RFC1123格式的时间戳
                now = datetime.now()
                date = format_date_time(mktime(now.timetuple()))

                # 拼接字符串
                signature_origin = "host: " + "ise-api.xfyun.cn" + "\n"
                signature_origin += "date: " + date + "\n"
                signature_origin += "GET " + "/v2/open-ise " + "HTTP/1.1"
                # 进行hmac-sha256进行加密
                signature_sha = hmac.new(self.APISecret.encode('utf-8'), signature_origin.encode('utf-8'),
                                         digestmod=hashlib.sha256).digest()
                signature_sha = base64.b64encode(signature_sha).decode(encoding='utf-8')

                authorization_origin = "api_key=\"%s\", algorithm=\"%s\", headers=\"%s\", signature=\"%s\"" % (
                    self.APIKey, "hmac-sha256", "host date request-line", signature_sha)
                authorization = base64.b64encode(authorization_origin.encode('utf-8')).decode(encoding='utf-8')
                # 将请求的鉴权参数组合为字典
                v = {
                    "authorization": authorization,
                    "date": date,
                    "host": "ise-api.xfyun.cn"
                }
                # 拼接鉴权参数，生成url
                url = url + '?' + urlencode(v)

                # 此处打印出建立连接时候的url,参考本demo的时候，比对相同参数时生成的url与自己代码生成的url是否一致
                print("date: ", date)
                print("v: ", v)
                print('websocket url :', url)
                return url

        # 收到websocket消息的处理
        def on_message(ws, message):
            try:
                code = json.loads(message)["code"]
                sid = json.loads(message)["sid"]
                if code != 0:
                    errMsg = json.loads(message)["message"]
                    print("sid:%s call error:%s code is:%s" % (sid, errMsg, code))

                else:
                    data = json.loads(message)["data"]
                    status = data["status"]
                    result = data["data"]
                    if (status == 2):
                        xml = base64.b64decode(result)
                        # python在windows上默认用gbk编码，print时需要做编码转换，mac等其他系统自行调整编码
                        print(xml.decode("gbk"))
                        xml_data = xml.decode("gbk")
                        # 解析XML数据
                        root = ET.fromstring(xml_data)
                        # 查找<read_sentence>标签
                        read_sentence = root.find('.//read_sentence')
                        if read_sentence is not None:
                            # 提取<read_sentence>中的属性
                            read_sentence_dict = {attr: read_sentence.attrib[attr] for attr in read_sentence.attrib}
                            # print("Read Sentence Attributes:", read_sentence_dict)
                        # 查找<sentence>标签
                        sentence = root.find('.//sentence')
                        if sentence is not None:
                            # 提取<sentence>中的属性
                            sentence_dict = {attr: sentence.attrib[attr] for attr in sentence.attrib}
                            # print("Sentence Attributes:", sentence_dict)
                            global gendu_score
                            gendu_score = sentence_dict
                            # print('genduscore::',gendu_score)
            except Exception as e:
                print("receive msg,but parse exception:", e)

        # 收到websocket错误的处理
        def on_error(ws, error):
            print("### error:", error)

        # 收到websocket关闭的处理
        def on_close(ws, close_status_code, close_msg):
            print("### closed ###")

        # 收到websocket连接建立的处理
        def on_open(ws):
            def run(*args):
                frameSize = 1280  # 每一帧的音频大小
                intervel = 0.04  # 发送音频间隔(单位:s)
                status = STATUS_FIRST_FRAME  # 音频的状态信息，标识音频是第一帧，还是中间帧、最后一帧

                with open(wsParam.AudioFile, "rb") as fp:
                    while True:
                        buf = fp.read(frameSize)
                        # 文件结束
                        if not buf:
                            status = STATUS_LAST_FRAME
                        # 第一帧处理
                        # 发送第一帧音频，带business 参数
                        # appid 必须带上，只需第一帧发送
                        if status == STATUS_FIRST_FRAME:
                            d = {"common": wsParam.CommonArgs,
                                 "business": wsParam.BusinessArgs,
                                 "data": {"status": 0}}
                            d = json.dumps(d)
                            ws.send(d)
                            status = STATUS_CONTINUE_FRAME
                        # 中间帧处理
                        elif status == STATUS_CONTINUE_FRAME:
                            d = {"business": {"cmd": "auw", "aus": 2, "aue": "raw"},
                                 "data": {"status": 1, "data": str(base64.b64encode(buf).decode())}}
                            ws.send(json.dumps(d))
                        # 最后一帧处理
                        elif status == STATUS_LAST_FRAME:
                            d = {"business": {"cmd": "auw", "aus": 4, "aue": "raw"},
                                 "data": {"status": 2, "data": str(base64.b64encode(buf).decode())}}
                            ws.send(json.dumps(d))
                            time.sleep(1)
                            break
                        # 模拟音频采样间隔
                        time.sleep(intervel)
                ws.close()

            thread.start_new_thread(run, ())

        if (1):

            time1 = datetime.now()
            # APPID、APISecret、APIKey信息在控制台——语音评测了（流式版）——服务接口认证信息处即可获取
            wsParam = Ws_Param(APPID='4626a4ce', APISecret='OTg0Y2E3MTIwOWU4MDc1YzQ1ZTE5MTgz',
                               APIKey='f7a10584a7234e3de1cfcf49f76b1054',
                               # AudioFile='E:/微信文件/WeChat Files/wxid_dkeyvqjvmg6222/FileStorage/File/2024-05/python065在线自主评测系统/程序/online_exam/media/kouyugendu/recording.wav',
                               AudioFile=(save_dir/'recording.wav'),
                               Text=TEXT)
            websocket.enableTrace(False)
            wsUrl = wsParam.create_url()
            ws = websocket.WebSocketApp(wsUrl, on_message=on_message, on_error=on_error, on_close=on_close)
            ws.on_open = on_open
            ws.run_forever(sslopt={"cert_reqs": ssl.CERT_NONE})
            time2 = datetime.now()
            print(time2 - time1)
        print('跟读评分：', gendu_score)
        accuracy_score = gendu_score['accuracy_score']
        fluency_score = gendu_score['fluency_score']
        standard_score = gendu_score['standard_score']
        total_score = gendu_score['total_score']
        if float(accuracy_score) <= 0.111:
            accuracy_score_pingjia = ('在你的口语中，有一些特定音素的发音存在偏差。'
                                      '例如‘th’音和‘r’音，这些发音上的不准确会影响到整体的清晰度和可理解性。建议你通过专业的发音训练课程，或是模仿标准发音的音频资料，'
                                      '专注于这些难点音素的练习。每天安排一定时间的针对性训练，结合舌位、唇形和呼吸控制的指导，将有助于你逐渐纠正这些发音问题，'
                                      '进而提升整体的发音质量。尽管目前您的口语流畅度有待提高，但我注意到您在努力尝试，这是非常值得鼓励的。')
        elif float(accuracy_score) > 0.111 and float(accuracy_score) < 0.155:
            accuracy_score_pingjia = ('在对话中，你偶尔会犯一些常见的语法错误，如时态误用、冠词不当以及主谓不一致等。'
                                      '这些问题可能源于对某些语法规则的理解不深或是在紧张情况下容易遗忘正确的用法。'
                                      '我建议你定期复习语法书籍中相关章节的内容，同时在日常练习中，有意识地检验自己是否正确应用了所学的语法规则。'
                                      '此外，通过写作练习，如日记或短文，可以加深对语法知识的记忆，并在实际应用中巩固这些规则。'
                                      '您在某些部分表现得相当不错，随着练习的增加，相信您的口语流畅度会有显著提升。')
        elif float(accuracy_score) > 0.1 and float(accuracy_score) <= 0.29:
            accuracy_score_pingjia = ('在你的口语中，语调较为单调，缺乏变化，这可能影响到表达的情感色彩和吸引听众的能力。语调的适当升降和停顿对于传达情绪和强调重点至关重要。'
                                      '建议你观察和模仿优秀的英语演讲者或播音员，注意他们如何运用语调来增强表达效果。练习朗读诗歌或戏剧剧本，也是提高语调变化的好方法。')
        elif float(accuracy_score) > 0.29 and float(accuracy_score) <= 0.5:
            accuracy_score_pingjia = ('你的流利度有所下降，有时会出现明显的停顿和重复。这可能是由于对句子结构的不确定性或对词汇的搜索。为了提高流利度，你可以通过大量听读英语材料，特别是那些带有自然对话节奏的音频，来培养语感。'
                                      '此外，练习即兴演讲和参与英语辩论活动，可以帮助你在压力下更好地组织语言和维持流畅的表达。')
        else:
            accuracy_score_pingjia = ('我发现你倾向于使用一些常见且基础的词汇，而较少尝试更高级或更具表现力的词汇。这可能会限制你表达思想的深度和细腻度。我建议你通过广泛阅读，尤其是英文原著、高质量的文章和学术论文，来丰富你的词汇库。'
                                      '同时，尝试在日常对话中主动使用新学到的词汇，这样不仅能增强语言的丰富性，还能提高你的口语表达能力。')
        if float(fluency_score) > 0.25: fluency_score_pingjia = ('你有时会表现出对对方话语理解的延迟，这可能导致回应不及时或答非所问的情况。有效沟通不仅在于说，也在于听。提升听力理解能力需要通过大量听英语材料，尤其是对话和讲座，来锻炼即时处理信息的能力。'
                                                                 '尝试边听边做笔记或复述听到的内容，这能帮助你更快地捕捉和处理信息，从而提高即时反应的速度和准确性。')
        if float(standard_score) > 0.002: standard_score_pingjia = ('在构建复杂的论点或叙述时，你的论述有时缺乏逻辑连贯性，导致听众难以跟随你的思路。逻辑清晰的表达需要使用合适的连接词和过渡句来确保段落之间的顺畅衔接。练习写作，尤其是议论文和故事叙述，可以帮助你学会如何构建有条理的论点。'
                                                                    '同时，在口语练习中，有意识地使用像‘firstly’, ‘moreover’, ‘on the other hand’这样的过渡词，可以大大提高你的表达逻辑性和连贯性。')
        if float(total_score) > 0.1: total_score_pingjia = total_score

        evaluation = accuracy_score_pingjia + '  ' + fluency_score_pingjia + '  ' + standard_score_pingjia + '  ' + total_score_pingjia
        return JsonResponse({'comment': evaluation})
    return JsonResponse({'comment': '请求方法错误'}, status=400)


# =======================================语音聊天======================================================

global_input_language = ''
global_input_language_dialect = ''
global_input_url = ''
global_xiaoyuzhong = 0
global_vcn = ''
chat_mode_code = 0


def set_chat_mode(request):
    global chat_mode_code
    if request.method == 'POST':
        mode = request.POST.get('mode')
        if mode is not None:
            try:
                chat_mode_code = int(mode)
                # 这里你可以做进一步的处理，比如保存到数据库
                print('-------对话模式：',chat_mode_code)
                # ...
                return JsonResponse({'status': 'success', 'message': f'Mode set to {chat_mode_code}'})
            except ValueError:
                return JsonResponse({'status': 'error', 'message': 'Invalid mode value'}, status=400)
        else:
            return JsonResponse({'status': 'error', 'message': 'No mode provided'}, status=400)
    else:
        return JsonResponse({'status': 'error', 'message': 'Only POST requests are allowed'}, status=405)


def setUserLanguage(request):
    global global_input_language, global_input_language_dialect, global_input_url
    if request.method == 'POST':


        global_input_language = request.POST.get('language')
        global_input_language_dialect = request.POST.get('dialect', '')
        global_input_url = request.POST.get('input_url', '')
        print("-------用户语言：", global_input_language)
        print("-------用户口音：", global_input_language_dialect)
        print("----------URL：", global_input_url)
        return JsonResponse(
            {"user_language": global_input_language, "user_language_dialect": global_input_language_dialect,
             "input_language_url": global_input_url})
    else:
        return JsonResponse({"status": "error", "message": "Invalid request method."})


def setAiLanguage(request):
    global global_xiaoyuzhong, global_vcn
    global_xiaoyuzhong = request.POST.get('xiaoyuzhong')
    global_vcn = request.POST.get('vcn')

    if global_xiaoyuzhong == 'English':
        global_xiaoyuzhong = 0
        global_vcn = 'x4_enus_luna_assist'

    elif global_xiaoyuzhong == 'French':
        global_xiaoyuzhong = 1
        global_vcn = 'x2_FrRgM_Lisa'
    elif global_xiaoyuzhong == 'Japanese':
        global_xiaoyuzhong = 1
        global_vcn = 'qianhui'

    print('-------是否为小语种：', global_xiaoyuzhong)
    print('--------AI发音口音：', global_vcn)
    return JsonResponse({"isxiaoyuzhong": global_xiaoyuzhong, "vcn": global_vcn})


results = []  # 全局变量用于存储结果


def dialogue(request):
    global results
    results = []  # 清空results列表，确保每次运行时它是空的
    save_dir = Path(settings.MEDIA_ROOT) / 'audio'


    if request.method == 'POST':
        audio_file = request.FILES['audio']
        # 保存音频文件到本地
        # file_path = os.path.join('E:/ai_dialogue/user', f"{uuid.uuid4()}.wav")
        file_path = os.path.join(save_dir, f"{uuid.uuid4()}.wav")

        with open(file_path, 'wb') as f:
            for chunk in audio_file.chunks():
                f.write(chunk)

                # 使用FFmpeg转换音频格式
        output_pcm_path = os.path.splitext(file_path)[0] + '.pcm'  # 修改文件扩展名为.pcm
        ffmpeg_command = [
            "ffmpeg",
            "-i", file_path,  # 输入文件
            "-ar", "16000",  # 设置采样率为16000Hz
            "-ac", "1",  # 设置声道数为1 (单声道)
            "-f", "s16le",  # 设置格式为16位小端存储的PCM
            output_pcm_path  # 输出文件
        ]

        try:
            subprocess.run(ffmpeg_command, check=True)
            print(f"Audio converted to PCM at {output_pcm_path}")
        except subprocess.CalledProcessError as e:
            print(f"FFmpeg process failed with error: {e}")
        finally:
            # 可选：删除原始wav文件
            os.remove(file_path)
        # =====================接受到用户语音，用户语音转文字=====================

        STATUS_FIRST_FRAME = 0  # 第一帧的标识
        STATUS_CONTINUE_FRAME = 1  # 中间帧标识
        STATUS_LAST_FRAME = 2  # 最后一帧的标识

        class Ws_Param(object):
            # 初始化
            def __init__(self, APPID, APIKey, APISecret, AudioFile):
                self.APPID = APPID
                self.APIKey = APIKey
                self.APISecret = APISecret
                self.AudioFile = AudioFile

                # 公共参数(common)
                self.CommonArgs = {"app_id": self.APPID}
                # 业务参数(business)，更多个性化参数可在官网查看
                self.BusinessArgs = {"domain": "iat", "language": global_input_language,
                                     "accent": global_input_language_dialect, "vinfo": 1,
                                     "vad_eos": 10000}

            # 生成url
            def create_url(self):
                url = global_input_url
                # 生成RFC1123格式的时间戳
                now = datetime.now()
                date = format_date_time(mktime(now.timetuple()))

                # 拼接字符串
                signature_origin = "host: " + "ws-api.xfyun.cn" + "\n"
                signature_origin += "date: " + date + "\n"
                signature_origin += "GET " + "/v2/iat " + "HTTP/1.1"
                # 进行hmac-sha256进行加密
                signature_sha = hmac.new(self.APISecret.encode('utf-8'), signature_origin.encode('utf-8'),
                                         digestmod=hashlib.sha256).digest()
                signature_sha = base64.b64encode(signature_sha).decode(encoding='utf-8')

                authorization_origin = "api_key=\"%s\", algorithm=\"%s\", headers=\"%s\", signature=\"%s\"" % (
                    self.APIKey, "hmac-sha256", "host date request-line", signature_sha)
                authorization = base64.b64encode(authorization_origin.encode('utf-8')).decode(encoding='utf-8')
                # 将请求的鉴权参数组合为字典
                v = {
                    "authorization": authorization,
                    "date": date,
                    "host": "ws-api.xfyun.cn"
                }
                # 拼接鉴权参数，生成url
                url = url + '?' + urlencode(v)
                # print("date: ",date)
                # print("v: ",v)
                # 此处打印出建立连接时候的url,参考本demo的时候可取消上方打印的注释，比对相同参数时生成的url与自己代码生成的url是否一致
                # print('websocket url :', url)
                return url

        # 收到websocket消息的处理
        def on_message(ws, message):

            try:
                code = json.loads(message)["code"]
                sid = json.loads(message)["sid"]
                if code != 0:
                    errMsg = json.loads(message)["message"]
                    print("sid:%s call error:%s code is:%s" % (sid, errMsg, code))

                else:
                    data = json.loads(message)["data"]["result"]["ws"]
                    # print(json.loads(message))
                    result = ""
                    for segment in data:
                        for candidate in segment["cw"]:
                            result += candidate["w"]
                    print("sid:%s call success!,data is:'%s'" % (sid, result))
                    results.append(result)  # 将结果追加到全局列表中
                    print(result)

            except Exception as e:
                print("receive msg,but parse exception:", e)

        # 收到websocket错误的处理
        def on_error(ws, error):
            print("### error:", error)

        # 收到websocket关闭的处理
        def on_close(ws, a, b):
            print("### closed ###")

        # 收到websocket连接建立的处理
        def on_open(ws):
            def run(*args):
                frameSize = 8000  # 每一帧的音频大小
                intervel = 0.04  # 发送音频间隔(单位:s)
                status = STATUS_FIRST_FRAME  # 音频的状态信息，标识音频是第一帧，还是中间帧、最后一帧

                with open(wsParam.AudioFile, "rb") as fp:
                    while True:
                        buf = fp.read(frameSize)
                        # 文件结束
                        if not buf:
                            status = STATUS_LAST_FRAME
                        # 第一帧处理
                        # 发送第一帧音频，带business 参数
                        # appid 必须带上，只需第一帧发送
                        if status == STATUS_FIRST_FRAME:

                            d = {"common": wsParam.CommonArgs,
                                 "business": wsParam.BusinessArgs,
                                 "data": {"status": 0, "format": "audio/L16;rate=16000",
                                          "audio": str(base64.b64encode(buf), 'utf-8'),
                                          "encoding": "raw"}}
                            d = json.dumps(d)
                            ws.send(d)
                            status = STATUS_CONTINUE_FRAME
                        # 中间帧处理
                        elif status == STATUS_CONTINUE_FRAME:
                            d = {"data": {"status": 1, "format": "audio/L16;rate=16000",
                                          "audio": str(base64.b64encode(buf), 'utf-8'),
                                          "encoding": "raw"}}
                            ws.send(json.dumps(d))
                        # 最后一帧处理
                        elif status == STATUS_LAST_FRAME:
                            d = {"data": {"status": 2, "format": "audio/L16;rate=16000",
                                          "audio": str(base64.b64encode(buf), 'utf-8'),
                                          "encoding": "raw"}}
                            ws.send(json.dumps(d))
                            time.sleep(1)
                            break
                        # 模拟音频采样间隔
                        time.sleep(intervel)
                ws.close()

            thread.start_new_thread(run, ())

        if (1):
            print('语音转文字')
            # course_list(request)
            time1 = datetime.now()
            # 这个地方是调用语音合成的接口，把文字转化为语音，并保存到本地，然后再调用websocket接口进行语音识别
            wsParam = Ws_Param(APPID='4626a4ce', APISecret='OTg0Y2E3MTIwOWU4MDc1YzQ1ZTE5MTgz',
                               APIKey='f7a10584a7234e3de1cfcf49f76b1054',
                               # AudioFile=r'E:/ai_dialogue/user/iat_pcm_16k.pcm'
                               AudioFile=output_pcm_path
                               )
            websocket.enableTrace(False)
            wsUrl = wsParam.create_url()
            ws = websocket.WebSocketApp(wsUrl, on_message=on_message, on_error=on_error, on_close=on_close)
            ws.on_open = on_open
            ws.run_forever(sslopt={"cert_reqs": ssl.CERT_NONE})
            time2 = datetime.now()
            print(time2 - time1)
            user_words = ''
            for item in results:
                user_words = user_words + item

            print('用户语音文字：', user_words)  # 好，此时已经获得用户的语音文字
            # =============================接下去把拿到的语音文字调用ai获得文字回答=================================================
            print("zhixngdaozheerle ", chat_mode_code)
            if chat_mode_code == 1:
                messages = [ChatMessage(
                    role="user",
                    content=  user_words
                )]
                handler = ChunkPrintHandler()
                a = spark.generate([messages], callbacks=[handler])
                first_reply_content = a.generations[0][0].message.content
                print('AI的文字回答：', first_reply_content)  # 好，此时已经获得了AI的文字回答，接下去把文字转化为语音
            elif chat_mode_code == 0:
                messages = [ChatMessage(
                    role="user",
                    content=user_words
                )]
                handler = ChunkPrintHandler()
                a = spark.generate([messages], callbacks=[handler])
                first_reply_content = a.generations[0][0].message.content
                print('AI的文字回答：', first_reply_content)  # 好，此时已经获得了AI的文字回答，接下去把文字转化为语音
 # ==================================接下去把拿到的AI文字调用语音合成获得音频回答========================================================================
            print('AI的文字回答：', first_reply_content)
            STATUS_FIRST_FRAME = 0  # 第一帧的标识
            STATUS_CONTINUE_FRAME = 1  # 中间帧标识
            STATUS_LAST_FRAME = 2  # 最后一帧的标识
            print('qqqqqq',global_vcn)
            class Ws_Param(object):
                # 初始化
                def __init__(self, APPID, APIKey, APISecret, Text):
                    self.APPID = APPID
                    self.APIKey = APIKey
                    self.APISecret = APISecret
                    self.Text = Text

                    # 公共参数(common)
                    self.CommonArgs = {"app_id": self.APPID}
                    # 业务参数(business)，更多个性化参数可在官网查看
                    self.BusinessArgs = {"aue": "raw", "auf": "audio/L16;rate=16000", "vcn": global_vcn, "tte": "utf8"}
                    if global_xiaoyuzhong:
                        # 使用小语种须使用以下方式，此处的unicode指的是 utf16小端的编码方式，即"UTF-16LE"”
                        self.Data = {"status": 2, "text": str(base64.b64encode(self.Text.encode('utf-16')), "UTF8")}
                    self.Data = {"status": 2, "text": str(base64.b64encode(self.Text.encode('utf-8')), "UTF8")}

                # 生成url
                def create_url(self):
                    url = 'wss://tts-api.xfyun.cn/v2/tts'
                    # 生成RFC1123格式的时间戳
                    now = datetime.now()
                    date = format_date_time(mktime(now.timetuple()))

                    # 拼接字符串
                    signature_origin = "host: " + "ws-api.xfyun.cn" + "\n"
                    signature_origin += "date: " + date + "\n"
                    signature_origin += "GET " + "/v2/tts " + "HTTP/1.1"
                    # 进行hmac-sha256进行加密
                    signature_sha = hmac.new(self.APISecret.encode('utf-8'), signature_origin.encode('utf-8'),
                                             digestmod=hashlib.sha256).digest()
                    signature_sha = base64.b64encode(signature_sha).decode(encoding='utf-8')

                    authorization_origin = "api_key=\"%s\", algorithm=\"%s\", headers=\"%s\", signature=\"%s\"" % (
                        self.APIKey, "hmac-sha256", "host date request-line", signature_sha)
                    authorization = base64.b64encode(authorization_origin.encode('utf-8')).decode(encoding='utf-8')
                    # 将请求的鉴权参数组合为字典
                    v = {
                        "authorization": authorization,
                        "date": date,
                        "host": "ws-api.xfyun.cn"
                    }
                    # 拼接鉴权参数，生成url
                    url = url + '?' + urlencode(v)
                    # print("date: ",date)
                    # print("v: ",v)
                    # 此处打印出建立连接时候的url,参考本demo的时候可取消上方打印的注释，比对相同参数时生成的url与自己代码生成的url是否一致
                    # print('websocket url :', url)
                    return url

            def on_message(ws, message):
                try:
                    message = json.loads(message)
                    code = message["code"]
                    sid = message["sid"]
                    audio = message["data"]["audio"]
                    audio = base64.b64decode(audio)
                    status = message["data"]["status"]
                    print(message)
                    if status == 2:
                        print("ws is closed")
                        ws.close()
                    if code != 0:
                        errMsg = message["message"]
                        print("sid:%s call error:%s code is:%s" % (sid, errMsg, code))
                    else:
                        save_dir = Path(settings.MEDIA_ROOT) / 'ai_reply_audio'
                        save_dir.mkdir(parents=True, exist_ok=True)
                        # 构建完整的文件路径
                        file_path = save_dir / 'demo.pcm'
                        # 写入音频数据
                        with file_path.open(mode='ab') as f:
                            f.write(audio)
                        # with open('E:/yuyinduihua/demo.pcm', 'ab') as f:
                        #     f.write(audio)


                except Exception as e:
                    print("receive msg,but parse exception:", e)

            # 收到websocket错误的处理
            def on_error(ws, error):
                print("### error:", error)

            # 收到websocket关闭的处理
            def on_close(ws, close_status_code, close_msg):
                print("### closed ###")

            # 收到websocket连接建立的处理
            def on_open(ws):
                def run(*args):
                    d = {"common": wsParam.CommonArgs,
                         "business": wsParam.BusinessArgs,
                         "data": wsParam.Data,
                         }
                    d = json.dumps(d)
                    print("------>开始发送文本数据")
                    ws.send(d)
                    # if os.path.exists('E:/yuyinduihua/demo.pcm'):
                    #     os.remove('E:/yuyinduihua/demo.pcm')
                    save_dir = Path(settings.MEDIA_ROOT) / 'ai_reply_audio'

                    # 确保目录存在
                    save_dir.mkdir(parents=True, exist_ok=True)

                    # 构建完整的文件路径
                    file_path = save_dir / 'demo.pcm'     #此时调用语音合成接口，并把pcm格式的音频文件保存到本地



                    # 检查文件是否存在，并删除
                    if file_path.exists():
                        file_path.unlink()

                thread.start_new_thread(run, ())

            if (1):
                # 测试时候在此处正确填写相关信息即可运行，这个是用于合成人工语言的。

                # 把ai的文字变成pcm格式
                # wsParam = Ws_Param(APPID='0df2556e', APISecret='NWQ3YzA1ODYxNDM5YTMxZjJjOGRkZjEz',
                #                    APIKey='0f5b68a7345b71e3cbff62aa34d73239',
                #                    Text=first_reply_content)
                wsParam = Ws_Param(APPID='b76aaa42', APISecret='MTZiZDZlYTZiNmNlZDJiNGU5MjQwYTMy',
                                   APIKey='42e19203f0397771ce0ef1491a982c2e',
                                   Text=first_reply_content)
                websocket.enableTrace(False)
                wsUrl = wsParam.create_url()
                ws = websocket.WebSocketApp(wsUrl, on_message=on_message, on_error=on_error, on_close=on_close)
                ws.on_open = on_open
                ws.run_forever(sslopt={"cert_reqs": ssl.CERT_NONE})
                # input_pcm_path = r"E:/yuyinduihua/demo.pcm"
                # output_wav_path = r"E:/yuyinduihua/demo.wav"
                # ffmpeg_exe_path = r"D:\tools\压缩包\ffmpeg-6.1.1-full_build-shared\bin\ffmpeg.exe"

                # =====================================================================================================
                # input_pcm_path = r"E:/yuyinduihua/demo.pcm"
                # output_wav_path = r"E:/yuyinduihua/demo.wav"
                save_dir = Path(settings.MEDIA_ROOT) / 'ai_reply_audio'

                # 确保目录存在
                save_dir.mkdir(parents=True, exist_ok=True)

                # 构建输入文件路径
                input_pcm_path = save_dir / 'demo.pcm'

                # 构建输出文件路径
                output_wav_path = save_dir / 'demo.wav'
                # timestamp = str(int(time.time()))

                # 构建输出文件路径，使用时间戳作为文件名的一部分
                # output_wav_path = save_dir / f'demo_{timestamp}.wav'
                ffmpeg_exe_path = r"D:\tools\压缩包\ffmpeg-6.1.1-full_build-shared\bin\ffmpeg.exe"

                # 构造FFmpeg命令
                ffmpeg_command = [
                    ffmpeg_exe_path,  # 使用完整路径

                    "-f", "s16le",  # 输入格式为16位小端存储的PCM
                    "-ar", "16000",  # 输入采样率为16000Hz
                    "-ac", "1",  # 输入声道数为1（单声道）
                    '-y',    # 覆盖输出文件
                    "-i", input_pcm_path,  # 输入文件
                    output_wav_path  # 输出文件
                ]

                # 执行FFmpeg命令
                try:
                    subprocess.run(ffmpeg_command, check=True)
                    print(f"PCM file has been converted to WAV at {output_wav_path}")
                except subprocess.CalledProcessError as e:
                    print(f"FFmpeg process failed with error: {e}")

            # 机器人回复

            with open(output_wav_path, 'rb') as f:
                bot_reply_audio = f.read()
            bot_reply_audio_base64 = base64.b64encode(bot_reply_audio).decode()
            return JsonResponse({'replyAudio': bot_reply_audio_base64})
        return JsonResponse({'error': 'Invalid request'}, status=400)








# ============================================课程封面列表函数=返回的页面是AIcourse.html========================================================

def course_list(request):
    role = request.session["role"]
    # username = request.session['username']
    username = '周涛'


    courses = Course.objects.all()
    processed_courses = []
    for course in courses:
        # print(course.title, course.description, course.instructor, course.image_url, course.video_url,course.detail_page_url,course.chapter)

        processed_course = {
            'title': course.title,
            'description': course.description,
            'instructor': course.instructor,
            'image_url': course.image_url,
            'video_url': course.video_url,
            'detail_page_url': course.detail_page_url,
            'chapter': course.chapter,
            'username': username,

        }
        print(processed_course)
        processed_courses.append(processed_course)
    return render(request, 'AIcourse.html', {'courses': processed_courses, 'username': username})



# ============================================单门课程详情页函数=返回的页面类似 1.html，代表第一门课程的详情页========================================================


def course_detail(request, course_id):
    # username = request.session['username']
    username = '周涛'

    print("course",course_id)       #单个课程封面页1，点击“立即参加”即可进入课程列表
    template_name = f'{course_id}.html'
    return render(request, template_name,{'username': username})

# ============================================视频课详情页11，有公告，测试，考试，讨论区等，返回的页面是11.html========================================================

def shangke_list(request, course_id, chapter):  #视频课详情页11，有公告，测试，考试，讨论区等
    template_name = f'{chapter}.html'
    return render(request, template_name,)

# ================================第一个导航按钮”公告“============展示视频公告信息，返回的页面是gonggao_11.html========================================================

def gonggao(request, course_id, chapter, announcement_id):
    # username = request.session['username']
    username = '周涛'

    # 通过ID获取课程信息
    course = Course.objects.get(id=course_id)

    # 创建一个包含课程信息的字典
    processed_course = {
        'title': course.title,
        'description': course.description,
        'instructor': course.instructor,
        'image_url': course.image_url,
        'video_url': course.video_url,
        'detail_page_url': course.detail_page_url,
        'chapter': course.chapter,
        'gonggao_context': course.gonggao_context,

    }

    # 设置模板名称
    template_name = f'gonggao_{course_id * 11}.html'

    # 将单个课程字典传递给模板
    return render(request, template_name, {'course': processed_course,'username': username})

# ===================第二个导航按钮”课件“============展示视频各小视频的卡片，可以有点赞收藏评论，记录观看进度等数据，返回的页面是kejian_11.html========================================================
def kejian(request, course_id, chapter, announcement_id):
    global current_course_id
    current_course_id = course_id

    course = Course.objects.get(id=course_id)
    # username = request.session['username']
    username = '周涛'

    # 获取所有视频URL相关的字段
    video_fields = [field for field in Course._meta.fields if field.name.startswith('video_') and not field.name.endswith('_students')]
    video_name_fields = [field for field in Course._meta.fields if field.name.startswith('video') and field.name.endswith('_name')]
    video_date_fields = [field for field in Course._meta.fields if field.name.startswith('video') and field.name.endswith('_date')]

    # 收集所有视频URL、名称和日期
    video_urls = [getattr(course, field.name) for field in video_fields if getattr(course, field.name)]
    video_names = [getattr(course, field.name) for field in video_name_fields if getattr(course, field.name)]
    video_dates = [
        getattr(course, field.name).strftime("%Y/%m/%d") if getattr(course, field.name) is not None else '2024.6.27'
        for field in video_date_fields
    ]

    # 添加调试信息
    print('视频URL字段:', [field.name for field in video_fields])
    print('视频名称字段:', [field.name for field in video_name_fields])
    print('视频日期字段:', [field.name for field in video_date_fields])
    print('视频URL:', video_urls)
    print('视频名称:', video_names)
    print('视频日期:', video_dates)

    # 检查视频信息是否匹配
    if len(video_urls) != len(video_names) or len(video_urls) != len(video_dates):
        print('视频信息不匹配，请检查字段名称和数据')

    video_info = [
        {'url': url, 'name': name, 'date': date}
        for url, name, date in zip(video_urls, video_names, video_dates)
    ]
    processed_course = {
        'title': course.title,
        'description': course.description,
        'instructor': course.instructor,
        'image_url': course.image_url,
        'detail_page_url': course.detail_page_url,
        'chapter': course.chapter,
        'video_info': video_info,
    }
    print('用户名', username)
    template_name = f'kejian_{course_id * 11}.html'
    return render(request, template_name, {'course': processed_course, 'username': username})





# ================================第三个导航按钮”测验与作业“============展示视频课程所对应的习题，返回的页面是test_11.html========================================================


def test(request, course_id, chapter, announcement_id, pid): #获取测试题的函数，动态抽取题目的实现就在这完成
    username = request.session.get('username')
    role = request.session.get('role')
    shijuan = get_object_or_404(TestPaper, id=pid)

    exam_time = shijuan.time * 60
    xuanzeti = ast.literal_eval(shijuan.pid)[0:10]
    panduanti = ast.literal_eval(shijuan.pid)[10:15]
    # jiandati = ast.literal_eval(shijuan.pid)[15:16]

    xuanzeti_info = [Question.objects.filter(id=i).first() for i in xuanzeti]
    panduanti_info = [Question.objects.filter(id=j).first() for j in panduanti]
    # jiandati_info = [Question.objects.filter(id=k).first() for k in jiandati]

    testid = course_id * 11
    course = get_object_or_404(Course, id=course_id)

    context = {
        'username': username,
        'role': role,
        'shijuan': shijuan,
        'exam_time': exam_time,
        'xuanzeti_info': xuanzeti_info,
        'panduanti_info': panduanti_info,
        # 'jiandati_info': jiandati_info,
        'testid': testid,
        'course': course,
    }
    template_name = f'test_{testid}.html'

    return render(request, template_name, context)


def submit_test(request):
    username = request.session.get('username')
    if request.method == 'POST':
        data = json.loads(request.body)
        test_id = data.get('test_id')
        score = data.get('score')
        video_id = data.get('video_id')

        # 确保 video_id 存在
        if not video_id:
            return JsonResponse({'status': 'fail', 'error': 'video_id not found'}, status=400)

        # 准备要保存的数据
        test_key = f'test_{test_id}'
        new_test_data = {
            test_key: {
                'score': score,
                'test_id': test_id
            },
            'student': username
        }

        # 更新或创建对应 video_id 的 course 记录
        course = get_object_or_404(Course, id=video_id)
        test_scores = course.test_score if course.test_score else []

        # 检查是否已有该学生的数据
        student_found = False
        for student_data in test_scores:
            if student_data['student'] == username:
                student_data[test_key] = new_test_data[test_key]  # 更新该学生的特定测试数据
                student_found = True
                break

        if not student_found:
            test_scores.append(new_test_data)  # 添加新的学生数据

        course.test_score = test_scores
        course.save()

        return JsonResponse({'status': 'success', 'score': score})
    else:
        return JsonResponse({'status': 'fail'}, status=400)

# ================================第四个导航按钮”考试“============展示视频考试信息，返回的页面是exam_11.html========================================================

# def exam(request,course_id,chapter,announcement_id):
#     examid = course_id * 11
#     course = Course.objects.get(id=course_id)
#     processed_course = {
#         'title': course.title,
#         'description': course.description,
#         'instructor': course.instructor,
#         'image_url': course.image_url,
#         'video_url': course.video_url,
#         'detail_page_url': course.detail_page_url,
#         'chapter': course.chapter,
#     }
#     add_knowledge_point("机器学习", "线性回归是一种监督学习方法...", 1)
#     template_name = f'exam_{examid}.html'
#     return render(request, template_name, {'course': processed_course})

# ================================================讨论================================================================
# def discuss(request, course_id, chapter, announcement_id):
#     examid = course_id * 11
#     course = Course.objects.all()
#     processed_course = {
#         'title': course.title,
#         'description': course.description,
#         'instructor': course.instructor,
#         'image_url': course.image_url,
#         'video_url': course.video_url,
#         'detail_page_url': course.detail_page_url,
#         'chapter': course.chapter,
#     }
#
#     template_name = f'discuss_{examid}.html'
#     return render(request, template_name, {'course': processed_course})




def discuss(request, course_id, chapter, announcement_id):
    username = request.session.get('username')
    number = request.session.get('number')


    examid = course_id * 11
    courses = Course.objects.all()

    processed_courses = [
        {
            'id': course.id,
            'title': course.title,
            'description': course.description,
            'instructor': course.instructor,
            'image_url': course.image_url,
            'video_url': course.video_url,
            'detail_page_url': course.detail_page_url,
            'chapter': course.chapter,
            'username': username,
            'number': number,
        }
        for course in courses
    ]

    template_name = f'discuss_{examid}.html'
    return render(request, template_name, {'courses': processed_courses})


# ============================更新学生的数据，点赞收藏评论，观看视频时长=============================================================

def update_student_data(request):

    username = request.session.get('username')
    if request.method == 'POST':
        try:
            data_str = request.body.decode('utf-8')
            data = json.loads(data_str)

            video_id = data.get('video_id')
            student_data = {
                'video_id': video_id,
                'liked': data.get('liked'),
                'favorited': data.get('favorited'),
                'comment': data.get('comment'),
                'progress': data.get('progress'),  # 获取播放进度
                'student': username,
            }
            try:
                course = Course.objects.get(pk=current_course_id)  # 假设你知道你要更新的是哪个课程
                field_name = f"video_{video_id}_students"
                students_list = getattr(course, field_name, None)  # 获取现有列表，如果没有则初始化为 None
                if students_list is None:
                    setattr(course, field_name, [])
                students_list = getattr(course, field_name)  # 再次获取列表，确保它已初始化

                if students_list is not None:  # 避免 None 类型的迭代
                    for i, student in enumerate(students_list):
                        if student['student'] == student_data['student']:
                            students_list[i] = student_data  # 替换相同的学生数据
                            break
                    else:
                        students_list.append(student_data)  # 如果没有找到相同的 student，添加新的学生数据

                setattr(course, field_name, students_list)  # 更新字段
                course.save()  # 保存更改

                return JsonResponse({"message": "Data saved successfully."})
            except ObjectDoesNotExist:
                return JsonResponse({"error": "Course not found"}, status=404)
        except json.JSONDecodeError as e:
            print("Error decoding JSON:", str(e))
            return JsonResponse({"error": "Invalid JSON data"}, status=400)


# =============================================无效但删除有风险代码=======================================================================
# =============================================无效但删除有风险代码=======================================================================

def submit_paper(request):
    if request.method == 'POST':
        # 处理提交的数据
        return JsonResponse({'status': 'success'})
    return JsonResponse({'status': 'failed'}, status=400)

def success_page(request):
    return render(request, 'success.html')
# =============================================无效但删除有风险代码=======================================================================
# =============================================无效但删除有风险代码=======================================================================











# ==============================================下面是VR===========================================================================

def japanese_museum(request):
    # 这里可以处理逻辑，如记录点击等
    return redirect('https://mshare.vrnewg.com/?nz30+Y0PuyQiht9Tvis0QmgjPwX4B1xJCC/ZVPejJvk=')  # 跳转到外部URL

def nanking_massacre_museum(request):
    return redirect('http://example.com/nanjing_massacre_museum_url')

def physics_lab(request):
    return redirect('http://example.com/physics_lab_url')
# =================================================获取学生各个数据代码===========================================================================

def get_student_data(username, course_id):
    student_data = []
    courses = Course.objects.filter(id=course_id)  # 获取指定课程

    for course in courses:
        course_data = {
            'course_name': course.title,
            'videos': [],
            'test_scores': [{'test_id': str(i), 'score': '未考试'} for i in range(1, 6)],  # 初始化5个固定测试
            'rowspan': 17  # 固定为17行：1行课程名+16行视频数据
        }

        # 提取视频观看数据
        for i in range(1, 17):  # 假设你有video_1_students到video_16_students
            field_name = f'video_{i}_students'
            students_list = getattr(course, field_name, None)
            video_data = {
                'video_id': i,
                'liked': '点赞未录入',
                'favorited': '收藏未录入',
                'progress': '进度未录入',
                'comment': '评论未录入',
            }


            if students_list:
                for student in students_list:
                    if student['student'] == username:
                        video_data = student
            course_data['videos'].append(video_data)

        # 提取测试成绩数据
        test_scores = course.test_score if course.test_score else []
        for score_data in test_scores:
            if score_data['student'] == username:
                for test_key, test_info in score_data.items():
                    if test_key != 'student':
                        for test in course_data['test_scores']:
                            if test['test_id'] == test_info['test_id']:
                                test['score'] = test_info['score']

        student_data.append(course_data)

    return student_data




def student_data_view(request):
    # username = request.session.get('username')
    username = '周涛'
    if username:
        student_topics = Topic.objects.filter(student_name=username)

        # 统计每个课程中的话题数量
        topic_counts = {}
        for topic in student_topics:
            course_title = topic.course.title
            if course_title in topic_counts:
                topic_counts[course_title] += 1
            else:
                topic_counts[course_title] = 1
        # for course_title in topic_counts:
        #     topic_counts[course_title] += 1


    if not username:
        return JsonResponse({'status': 'fail', 'error': 'user not logged in'}, status=400)

    course_id = request.GET.get('course_id', 1)  # 默认显示课程ID为1
    student_data = get_student_data(username, course_id)
    all_courses = Course.objects.all()  # 获取所有课程列表

    # 初始化变量
    total_likes = 0
    total_favorites = 0
    total_score = 0
    comment_dict = {}
    sentiment_scores = {}  # 存储评论的情感得分
    progress_list = [0] * 16  # 存储16个视频的播放进度
    useful_progress_number = 0  # 播放进度大于30的个数
    test_scores = [0] * 5  # 存储5门测试的具体分数
    test_participation_count = 0  # 参与考试的数量
    num_videos = 16
    num_tests = 5

    # 计算点赞量、收藏量和评论
    for video in student_data[0]['videos']:
        video_id = video.get('video_id', 0) - 1  # Assuming video_id starts from 1 and maps to index 0-15
        liked = video.get('liked', 0)
        favorited = video.get('favorited', 0)
        progress = video.get('progress', 0)
        comment = video.get('comment', '评论未录入')

        # 计算总点赞量和收藏量
        if liked == 1:
            total_likes += 1
        if favorited == 1:
            total_favorites += 1

        # 更新视频播放进度列表
        if 0 <= video_id < num_videos:
            progress_list[video_id] = progress

        # 计算进度大于30的个数
        if isinstance(progress, int) and progress > 30:
            useful_progress_number += 1

        # 收集评论
        comment_dict[video['video_id']] = comment
        print('pl',comment)

        # # 情感分析并存储得分
        # sentiment_score = HanLP.sentiment_analysis(comment)
        # sentiment_scores[video['video_id']] = sentiment_score

    # 计算测试卷分数和参与考试数量
    for i, test in enumerate(student_data[0]['test_scores']):
        score = test.get('score', 0)
        if isinstance(score, int):
            total_score += score
            test_scores[i] = score
            test_participation_count += 1

    # 计算点赞百分比和收藏百分比
    like_percentage = total_likes / num_videos
    favorite_percentage = total_favorites / num_videos

    # 计算测试卷平均分
    average_score = total_score / num_tests

    # 计算测试参与度
    participation_rate = test_participation_count / num_tests

    # 打印变量
    print('学生姓名:', username)
    print('总点赞量:', total_likes)
    print('点赞比例:', like_percentage)
    print('总收藏量:', total_favorites)
    print('收藏比例:', favorite_percentage)
    print('测试分数列表:', test_scores)
    print('测试平均分:', average_score)
    print('测试参与度:', participation_rate)
    print('评论集:', comment_dict)
    print('评论情感得分:', sentiment_scores)
    print('播放进度列表:', progress_list)
    print('有效观看视频数量:', useful_progress_number)
    print('发布的课程话题数量', topic_counts)


    return render(request, 'student_data.html', {
        'student_data': student_data,
        'all_courses': all_courses,
        'course': Course.objects.get(id=course_id),
        'username': username,
        'total_likes': total_likes,
        'like_percentage': like_percentage,
        'total_favorites': total_favorites,
        'favorite_percentage': favorite_percentage,
        'progress_list': progress_list,
        'useful_progress_number': useful_progress_number,
        'average_score': average_score,
        'test_scores': test_scores,
        'participation_rate': participation_rate,
        'comment_dict': comment_dict,
        'sentiment_scores': sentiment_scores,
        'topic_counts': topic_counts  # 将话题数量传递给前端
    })


# ===================================================================================================

def get_courses_by_module(request, module_id):
    courses = Course.objects.filter(module_id=module_id).values('id', 'title')
    return JsonResponse({'courses': list(courses)})



def api_courses(request):
    if request.method == 'GET':
        courses = Course.objects.all().values('id', 'name')
        return JsonResponse(list(courses), safe=False)


def api_topics(request):
    if request.method == 'POST':
        try:
            student_name = request.session.get('username')
            number = request.session.get('number')

            module = request.POST.get('module')
            course_id = request.POST.get('course_id')
            title = request.POST.get('title')
            description = request.POST.get('description')

            if not course_id:
                return JsonResponse({'status': 'error', 'message': 'Course ID is required.'}, status=400)



            topic = Topic(
                title=title,
                description=description,
                module=module,
                course_id=course_id,
                author_id=1,  # Assuming you have a default author or get it from session
                student_name=student_name,
                student_number=number
            )
            topic.save()



            return JsonResponse({'status': 'success'}, status=201)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=400)
    else:
        return JsonResponse({'status': 'error', 'message': 'Invalid request method.'}, status=405)

def api_topics_list(request):
    student_name = request.session.get('username')
    if request.method == 'GET':



        topics = Topic.objects.all().values('id', 'title', 'student_name', 'description', 'author__username', 'created_at', 'updated_at', 'course__title', 'module')
        return JsonResponse(list(topics), safe=False)
    else:
        return JsonResponse({'status': 'error', 'message': 'Invalid request method.'}, status=405)




# def get_topic_comments(request, topic_id):
#     # 根据topic_id从数据库获取评论
#     comments = Comment.objects.filter(topic_id=topic_id)
#     # 将评论转换为字典列表
#     comments_list = [
#         {'id': c.id, 'content': c.content, 'author_name': f"{c.author.first_name} {c.author.last_name}".strip(),
#          'created_at': c.created_at} for c in comments]
#     # 返回JSON响应
#     return JsonResponse({'comments': comments_list})





@csrf_exempt
def api_comments(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            content = data.get('content')
            topic_id = data.get('topic_id')
            parent_comment_id = data.get('parent_comment_id')

            # 获取当前登录的学生
            student_name = request.session.get('username')

            if not content or not topic_id:
                return JsonResponse({'status': 'error', 'message': 'Content and Topic ID are required.'}, status=400)

            # 创建Comment实例
            comment = Comment(
                content=content,
                author_id=1,  # 设置author_id为1
                topic_id=topic_id,
                parent_comment_id=parent_comment_id if parent_comment_id else None,  # 处理 parent_comment_id 为空的情况
                student_name=student_name
            )
            comment.save()

            # 返回成功响应，包含新评论的完整信息
            return JsonResponse({
                'status': 'success',
                'comment': {
                    'id': comment.id,
                    'content': comment.content,
                    'student_name': comment.student_name,
                    'created_at': comment.created_at,
                    # 其他评论信息...
                }
            }, status=201)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=400)
    else:
        return JsonResponse({'status': 'error', 'message': 'Invalid request method.'}, status=405)
def api_comments_list(request, topic_id):
    comments = Comment.objects.filter(topic_id=topic_id, parent_comment=None)
    comments_data = []
    for comment in comments:
        comment_data = {
            'id': comment.id,
            'content': comment.content,
            'student_name': comment.student_name,
            'created_at': comment.created_at,
            'topic_id': topic_id,
            'replies': []
        }
        replies = comment.replies.all()
        for reply in replies:
            reply_data = {
                'id': reply.id,
                'content': reply.content,
                'student_name': reply.student_name,
                'created_at': reply.created_at
            }
            comment_data['replies'].append(reply_data)
        comments_data.append(comment_data)
    return JsonResponse(comments_data, safe=False)

def topic_detail(request, topic_id):
    try:
        topic = Topic.objects.get(id=topic_id)
        response_data = {
            "id": topic.id,
            "module": topic.module,
            "title": topic.title,
            "description": topic.description,
            "student_name": topic.student_name,
            "created_at": topic.created_at,
            "updated_at": topic.updated_at,
        }
        return JsonResponse(response_data)
    except Topic.DoesNotExist:
        return JsonResponse({"error": "Topic not found"}, status=404)



def fetch_comments(request):
    topic_id = request.GET.get('topic_id')
    comments = Comment.objects.filter(topic_id=topic_id).values('id', 'content', 'author', 'created_at', 'parent_comment')
    comments_list = list(comments)

    for comment in comments_list:
        replies = Comment.objects.filter(parent_comment=comment['id']).values('id', 'content', 'author', 'created_at')
        comment['replies'] = list(replies)

    return JsonResponse({'comments': comments_list})


def get_detailed_datas(request):
    # 根据实际需求获取详细数据
    detailed_data = {
        'example_key': 'example_value',
        # 添加更多数据
    }
    return JsonResponse(detailed_data)


global course_id_to_analyse

def get_course_id_to_analyse(request):
    course_id_to_analyse = 1
    if request.method == 'POST':
        courseName = request.POST.get('courseName')

        if courseName == '智能机器人系统':
            course_id_to_analyse = 1
        elif courseName == '智能与信息社会':
            course_id_to_analyse = 2
        elif courseName == '人工智能原理':
            course_id_to_analyse = 3
        elif courseName == '云计算技术与应用':
            course_id_to_analyse = 4
        elif courseName == '深度学习及其应用':
            course_id_to_analyse = 5
        elif courseName == '大数据算法':
            course_id_to_analyse = 6
        elif courseName == '可持续智能城镇化':
            course_id_to_analyse = 7
        elif courseName == '模型与算法':
            course_id_to_analyse = 8
        elif courseName == '可视化导论':
            course_id_to_analyse = 9
        print('选择课程id：', course_id_to_analyse)

        return JsonResponse({'status': 'success', 'message': f'Received course ID: 1'})


def analyzeScores(request):
    latest_news = News.objects.order_by('-id')[:10]
    print('latest_news是。。。。',latest_news)

    course_id_to_analyse = 1
    # username = request.session.get('username')
    username = '周'
    if username:
        student_topics = Topic.objects.filter(student_name=username)
        # 统计每个课程中的话题数量
        topic_counts = {}
        for topic in student_topics:
            course_title = topic.course.title
            if course_title in topic_counts:
                topic_counts[course_title] += 1
            else:
                topic_counts[course_title] = 1
    if not username:
        return JsonResponse({'status': 'fail', 'error': 'user not logged in'}, status=400)
    course_id = request.GET.get('course_id', course_id_to_analyse)  # 默认显示课程ID为1
    student_data = get_student_data(username, course_id)
    print('当前查询课程id', course_id_to_analyse)
    all_courses = Course.objects.all()  # 获取所有课程列表
    # 初始化变量
    total_likes = 0  # 点赞量
    total_favorites = 0  # 收藏量
    total_score = 0  # 测试卷总分
    comment_dict = {}  # 存储评论内容
    sentiment_scores = {}  # 存储评论的情感得分
    progress_list = [0] * 16  # 存储16个视频的播放进度
    useful_progress_number = 0  # 播放进度大于30的个数
    test_scores = [0] * 5  # 存储5门测试的具体分数
    test_participation_count = 0  # 参与考试的数量
    num_videos = 16  # 视频数量
    num_tests = 5  # 测试卷数量
    # 计算点赞量、收藏量和评论
    for video in student_data[0]['videos']:
        video_id = video.get('video_id', 0) - 1  # Assuming video_id starts from 1 and maps to index 0-15
        liked = video.get('liked', 0)
        favorited = video.get('favorited', 0)
        progress = video.get('progress', 0)
        comment = video.get('comment', '评论未录入')
        # 计算总点赞量和收藏量
        if liked == 1:
            total_likes += 1
        if favorited == 1:
            total_favorites += 1
        # 更新视频播放进度列表
        if 0 <= video_id < num_videos:
            progress_list[video_id] = progress
        # 计算进度大于30的个数
        if isinstance(progress, int) and progress > 30:
            useful_progress_number += 1
        unuseful_progress_number = num_videos - useful_progress_number

        # 计算测试卷平均分
        average_score = total_score / num_tests
        # 计算测试参与度
        participation_rate = test_participation_count / num_tests
        # 收集评论
        comment_dict[video['video_id']] = comment
        # 情感分析并存储得分
        sentiment_score = HanLP.sentiment_analysis(comment)
        sentiment_scores[video['video_id']] = sentiment_score
        emo_variables = list(sentiment_scores.values())




    print('emo_variables是。。。。',emo_variables)
    # emo_variables = [0.11357155442237854, 0.21357155442237854, -0.11357155442237854, -0.11357155442237854, -0.11357155442237854, -0.11357155442237854, -0.11357155442237854, -0.11357155442237854, -0.11357155442237854, -0.11357155442237854, -0.11357155442237854, -0.11357155442237854, -0.11357155442237854, -0.11357155442237854, -0.11357155442237854, -0.11357155442237854]


    emo_1 = emo_variables[0] * 100
    emo_2 = emo_variables[1] * 100
    emo_3 = emo_variables[2] * 100
    emo_4 = emo_variables[3] * 100
    emo_5 = emo_variables[4] * 100
    emo_6 = emo_variables[5] * 100
    emo_7 = emo_variables[6] * 100
    emo_8 = emo_variables[7] * 100
    emo_9 = emo_variables[8] * 100
    emo_10 = emo_variables[9] * 100
    emo_11 = emo_variables[10] * 100
    emo_12 = emo_variables[11] * 100
    emo_13 = emo_variables[12] * 100
    emo_14 = emo_variables[13] * 100
    emo_15 = emo_variables[14] * 100
    emo_16 = emo_variables[15] * 100

    # emo_1 = 30
    # emo_2 = 40
    # emo_3 = 50
    # emo_4 = 60
    # emo_5 = 70
    # emo_6 = 80
    # emo_7 = 90
    # emo_8 = 10
    # emo_9 = 20
    # emo_10 = 30
    # emo_11 = 40
    # emo_12 = 50
    # emo_13 = 60
    # emo_14 = 70
    # emo_15 = 80
    # emo_16 = 90

    try:
        course = Course.objects.get(id=course_id)
    except Course.DoesNotExist:
        return JsonResponse({'status': 'fail', 'error': 'Course not found'}, status=404)

    low_emo_numbers = [2]
    for i, score in enumerate(emo_variables):
        if score * 100 < 40:
            # `emo_`后面的数字为i+1
            low_emo_numbers.append(i + 1)

    # 打印小于40的情感得分对应的`emo_`后面的数字
    print("当前课程低情感得分的视频号:")
    print(low_emo_numbers)

    low_knowledge_points = [course.knowledge[str(num)] for num in low_emo_numbers if str(num) in course.knowledge]

    # 打印或返回这些知识点
    print("低情感得分的知识点:")
    print(low_knowledge_points)
    # 在下面抽取低情感得分知识点的视频
    related_videos = IndividualVideo.objects.filter(knowledge__in=low_knowledge_points).values('title', 'url')

    # 将结果转换为一个列表，便于序列化到JSON
    video_list = list(related_videos)
    print("相关视频:")
    print(video_list)

    # 计算测试卷分数和参与考试数量
    for i, test in enumerate(student_data[0]['test_scores']):
        score = test.get('score', 0)
        if isinstance(score, int):
            total_score += score
            test_scores[i] = score
            test_participation_count += 1
    # 计算点赞百分比和收藏百分比
    like_percentage = total_likes / num_videos
    favorite_percentage = total_favorites / num_videos
    test_scores = student_data[0].get('test_scores', [{'score': 0}] * 5)  # 确保总是有5个字典，每个字典至少有'score'键

    # 提取分数
    test_score_1 = test_scores[0].get('score', 0)
    test_score_2 = test_scores[1].get('score', 0)
    test_score_3 = test_scores[2].get('score', 0)
    test_score_4 = test_scores[3].get('score', 0)
    test_score_5 = test_scores[4].get('score', 0)

    # 将非数字的'score'转换为0
    test_score_1 = int(test_score_1) if isinstance(test_score_1, (int, float)) else 0
    test_score_2 = int(test_score_2) if isinstance(test_score_2, (int, float)) else 0
    test_score_3 = int(test_score_3) if isinstance(test_score_3, (int, float)) else 0
    test_score_4 = int(test_score_4) if isinstance(test_score_4, (int, float)) else 0
    test_score_5 = int(test_score_5) if isinstance(test_score_5, (int, float)) else 0

    # 打印变量
    # print('学生姓名:', username)
    # print('总点赞量:', total_likes)
    # print('点赞比例:', like_percentage)
    # print('总收藏量:', total_favorites)
    # print('收藏比例:', favorite_percentage)
    # print('测试分数列表:', test_scores)
    # print('测试平均分:', average_score)
    # print('测试参与度:', participation_rate)
    # print('评论集:', comment_dict)
    # print('评论情感得分:', sentiment_scores)
    # print('播放进度列表:', progress_list)
    # print('有效观看视频数量:', useful_progress_number)
    # print('发布的课程话题数量', topic_counts)

    video_progresses = {}
    for index, progress in enumerate(progress_list):
        if progress == "进度未录入":
            progress_list[index] = 0  # 直接修改列表中的值

    # 现在，progress_list中的"进度未录入"已经被替换成0
    progress_1, progress_2, progress_3, progress_4, progress_5, progress_6, progress_7, progress_8, progress_9, progress_10, progress_11, progress_12, progress_13, progress_14, progress_15, progress_16 = progress_list[
                                                                                                                                                                                                            :16]
    # print('哈哈哈哈哈', progress_1, progress_2, progress_3, progress_4, progress_5, progress_6, progress_7, progress_8,
    #       progress_9, progress_10, progress_11, progress_12, progress_13, progress_14, progress_15, progress_16)
    # ===================================话题数量=============================================================

    topic_counts = {}
    for topic in student_topics:
        course_title = topic.course.title
        if course_title in topic_counts:
            topic_counts[course_title] += 1
        else:
            topic_counts[course_title] = 1

    topic_1 = topic_counts.get('智能机器人系统', 0)
    topic_2 = topic_counts.get('智能与信息社会', 0)
    topic_3 = topic_counts.get('人工智能原理', 0)
    topic_4 = topic_counts.get('云计算技术与应用', 0)
    topic_5 = topic_counts.get('深度学习及其应用', 0)
    topic_6 = topic_counts.get('大数据算法', 0)
    topic_7 = topic_counts.get('可持续智能城镇化', 0)
    topic_8 = topic_counts.get('模型与算法', 0)
    topic_9 = topic_counts.get('可视化导论', 0)

    # ==========================================考试成绩====================================================
    # recent_records = Record.objects.filter(name=username).order_by('-exam_time')[:5]
    #
    # # 初始化成绩列表
    # exam_test_scores = []
    #
    # # 遍历最近的记录，收集考试成绩
    # for record in recent_records:
    #     exam_test_scores.append({
    #         'exam_time': record.exam_time.strftime('%Y-%m-%d %H:%M:%S'),  # 格式化考试时间
    #         'score': record.grade,  # 考试成绩
    #         'test_paper_id': record.test_paper_id,  # 试卷ID
    #     })
    #
    # while len(exam_test_scores) < 5:
    #     exam_test_scores.append({
    #         'exam_time': None,
    #         'score': 0,
    #         'test_paper_id': None,
    #     })
    #
    # exam_1 = exam_test_scores[0]['score'] if len(exam_test_scores) > 0 else 0
    # exam_2 = exam_test_scores[1]['score'] if len(exam_test_scores) > 1 else 0
    # exam_3 = exam_test_scores[2]['score'] if len(exam_test_scores) > 2 else 0
    # exam_4 = exam_test_scores[3]['score'] if len(exam_test_scores) > 3 else 0
    # exam_5 = exam_test_scores[4]['score'] if len(exam_test_scores) > 4 else 0
    #
    # average_exam_score = (exam_1 + exam_2 + exam_3 + exam_4 + exam_5) / 5 if len(exam_test_scores) > 0 else 0



    # url = "http://tech.chinadaily.com.cn/5b8f760ea310030f813ed4c4"
    # response = requests.get(url)
    # response.encoding = "utf-8"
    # html = response.text
    # soup = BeautifulSoup(html, "lxml")
    #
    # divs = soup.find_all("div", class_="busBox3")
    #
    # for div in divs:
    #     title_tag = div.find("h3")
    #     date_tag = div.find('b')
    #
    #     if title_tag and date_tag:
    #         title = title_tag.text.strip()
    #         url = title_tag.find('a')['href']
    #         date = date_tag.text.strip()
    #         date_obj = datetime.strptime(date, '%Y-%m-%d %H:%M')
    #
    #         # 只保留日期部分
    #         date_only = date_obj.date()
    #
    #         # 现在可以安全地保存到Django的DateField中
    #         news = News(title=title, url=url, date=date_only)
    #         news.save()
    #         print(f'新闻标题: {title}, 新闻链接: {url}, 新闻日期: {date_only}')
    #
    #
    #
    #     print('最新新闻：ssss')
    #     latest_news = News.objects.order_by('-id')[:10]
    #     print(type('latest_newsd的类型是',latest_news))
    #     print(latest_news)

        # 检查文件是否存在，如果存在则加载现有工作簿，否则创建一个新的
    try:
        workbook = openpyxl.load_workbook('output_train_data.xlsx')
    except FileNotFoundError:
        workbook = openpyxl.Workbook()

        # 确保我们正在使用的表单是第一个活动表单
    worksheet = workbook.active

    # 如果工作表为空，写入列标题
    if worksheet.max_row == 0:
        headers = ['Username', 'Total Likes', 'Like Percentage', 'Total Favorites', 'Favorite Percentage',
                   'Useful Progress Number', 'Unuseful Progress Number', 'Average Score', 'Participation Rate',
                   'Test 1', 'Test 2', 'Test 3', 'Test 4', 'Test 5', 'Average Exam Score',
                   'Topic 1', 'Topic 2', 'Topic 3', 'Topic 4', 'Topic 5', 'Topic 6', 'Topic 7', 'Topic 8', 'Topic 9',
                   'Progress 1', 'Progress 2', 'Progress 3', 'Progress 4', 'Progress 5', 'Progress 6', 'Progress 7',
                   'Progress 8', 'Progress 9', 'Progress 10', 'Progress 11', 'Progress 12', 'Progress 13',
                   'Progress 14', 'Progress 15', 'Progress 16',
                   'Emo 1', 'Emo 2', 'Emo 3', 'Emo 4', 'Emo 5', 'Emo 6', 'Emo 7', 'Emo 8', 'Emo 9', 'Emo 10',
                   'Emo 11', 'Emo 12', 'Emo 13', 'Emo 14', 'Emo 15', 'Emo 16', 'Average_Exam_Score']
        worksheet.append(headers)

    # 写入数据行
    data_row = [
        username, total_likes, like_percentage, total_favorites, favorite_percentage,
        useful_progress_number, unuseful_progress_number, average_score, participation_rate,
        test_score_1, test_score_2, test_score_3, test_score_4, test_score_5,
        # average_exam_score,
        87,
        topic_1, topic_2, topic_3, topic_4, topic_5, topic_6, topic_7, topic_8, topic_9,
        progress_1, progress_2, progress_3, progress_4, progress_5, progress_6, progress_7,
        progress_8, progress_9, progress_10, progress_11, progress_12, progress_13,
        progress_14, progress_15, progress_16,
        emo_1, emo_2, emo_3, emo_4, emo_5, emo_6, emo_7, emo_8, emo_9, emo_10,
        emo_11, emo_12, emo_13, emo_14, emo_15, emo_16,
        # average_exam_score
        87
    ]

    worksheet.append(data_row)

    # 保存工作簿
    workbook.save('output_train_data.xlsx')
    # 加载数据
    # desktop_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
    # file_path = os.path.join(desktop_path, 'student_data.xlsx')
    # data = pd.read_excel(file_path)
    #
    # # 数据预处理
    # data = data.drop(columns=['Username'])
    # X = data.drop(columns=['Average Exam Score'])
    # y = data['Average Exam Score']

    # 立刻从保存数据的同一个文件加载数据
    file_path = 'output_train_data.xlsx'
    data = pd.read_excel(file_path)

    # 数据预处理
    data = data.drop(columns=['Username'])
    X = data.drop(columns=['Average Exam Score'])
    y = data['Average Exam Score']

    # 分割数据
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

    # 标准化数据
    scaler = StandardScaler()
    X_train = scaler.fit_transform(X_train)
    X_test = scaler.transform(X_test)

    # 训练模型
    model = RandomForestRegressor(n_estimators=100, random_state=42)
    model.fit(X_train, y_train)

    # 评估模型
    y_pred = model.predict(X_test)
    mae = mean_absolute_error(y_test, y_pred)
    mse = mean_squared_error(y_test, y_pred)
    rmse = mse ** 0.5

    # print(f'Mean Absolute Error: {mae}')
    # print(f'Mean Squared Error: {mse}')
    # print(f'Root Mean Squared Error: {rmse}')

    # 预测新数据函数
    def predict_new_data(new_data):
        new_data_scaled = scaler.transform(new_data)
        predictions = model.predict(new_data_scaled)
        return predictions

    # 示例：预测最新的学生数据（即最后一行）
    last_student_data = X.tail(1)  # 选择DataFrame的最后一行
    last_student_predictions = predict_new_data(last_student_data)
    print(f'当前学生预测得分: {last_student_predictions[0]}')
    global_predict_student_score = int(last_student_predictions[0])


    # 将测试分数传递给上下文
    context = {
        'latest_news': latest_news,  # 添加最新新闻到context
        'all_courses': all_courses,
        'test_scores': test_scores,
        'total_likes': total_likes,
        'like_percentage': like_percentage,
        'total_favorites': total_favorites,
        'favorite_percentage': favorite_percentage,
        'useful_progress_number': useful_progress_number,
        'unuseful_progress_number': unuseful_progress_number,
        'total_likes': total_likes,
        'total_favorites': total_favorites,
        'like_percentage': like_percentage,
        'favorite_percentage': favorite_percentage,
        'progress_1': progress_1,
        'progress_2': progress_2,
        'progress_3': progress_3,
        'progress_4': progress_4,
        'progress_5': progress_5,
        'progress_6': progress_6,
        'progress_7': progress_7,
        'progress_8': progress_8,
        'progress_9': progress_9,
        'progress_10': progress_10,

        'progress_11': progress_11,
        'progress_12': progress_12,
        'progress_13': progress_13,
        'progress_14': progress_14,
        'progress_15': progress_15,
        'progress_16': progress_16,

        'topic_1': topic_1,
        'topic_2': topic_2,
        'topic_3': topic_3,
        'topic_4': topic_4,
        'topic_5': topic_5,
        'topic_6': topic_6,
        'topic_7': topic_7,
        'topic_8': topic_8,
        'topic_9': topic_9,

        # 'exam_1': exam_1,
        # 'exam_2': exam_2,
        # 'exam_3': exam_3,
        # 'exam_4': exam_4,
        # 'exam_5': exam_5,

        'exam_1': 87,
        'exam_2': 82,
        'exam_3': 79,
        'exam_4': 83,
        'exam_5': 87,

        'predict_score': int(last_student_predictions[0]),
        'video_list': video_list,

        'emo_1': emo_1,
        'emo_2': emo_2,
        'emo_3': emo_3,
        'emo_4': emo_4,
        'emo_5': emo_5,
        'emo_6': emo_6,
        'emo_7': emo_7,
        'emo_8': emo_8,
        'emo_9': emo_9,
        'emo_10': emo_10,
        'emo_11': emo_11,
        'emo_12': emo_12,
        'emo_13': emo_13,
        'emo_14': emo_14,
        'emo_15': emo_15,
        'emo_16': emo_16,

    }

    return render(request, 'dapingmu.html', context)


def generate_exam(request):
    difficulty = request.GET.get('difficulty', 'A')  # 获取请求参数中的难度等级

    # 从数据库中抽取指定难度等级的试题
    timu = IndividualTimu.objects.filter(difficulty=difficulty).first()

    if not timu:
        return HttpResponse("No questions found for the specified difficulty.", status=404)

    # 创建Word文档
    doc = Document()

    # 设置文档的页面边距
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # 添加密封线
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = '------密--封--线--内--不要--答--题--------'
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(14)
    run.bold = True


    # 添加标题和学校年级姓名信息
    title = doc.add_heading('《人工智能基础检测卷》（个性化）', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('学校：____________________        年级：____________________    姓名：____________________')
    doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)

    # 各大题的评分表
    doc.add_paragraph('一、选择题：每题2分，共20分\n二、填空题：每题3分，共30分\n三、计算题：每题10分，共50分')


    # 添加选择题
    doc.add_heading('一、选择题', level=1)
    for i in range(1, 11):
        question = getattr(timu, f'choice_question{i}')
        doc.add_paragraph(f'{i}. {question["title"]}')
        doc.add_paragraph(f'A. {question["A"]}')
        doc.add_paragraph(f'B. {question["B"]}')
        doc.add_paragraph(f'C. {question["C"]}')
        doc.add_paragraph(f'D. {question["D"]}')
        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)


    # 添加填空题
    doc.add_heading('二、填空题', level=1)
    for i in range(1, 11):
        question = getattr(timu, f'tiankong{i}')
        doc.add_paragraph(f'{i}. {question["title"]}')



    # 添加计算题
    doc.add_heading('三、计算题', level=1)
    for i in range(1, 6):
        question = getattr(timu, f'jisuanti{i}')
        doc.add_paragraph(f'{i}. {question["title"]}')
        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)  # 添加答题空间
        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)

    # 保存文档并返回HttpResponse
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=personalized_exam.docx'
    doc.save(response)

    return response


def get_exam_info(request):
    difficulty = request.GET.get('difficulty', 'A')
    timu = IndividualTimu.objects.filter(difficulty=difficulty).first()

    if not timu:
        return JsonResponse({"error": "No questions found for the specified difficulty."}, status=404)

    return JsonResponse({"title": timu.title, "difficulty": timu.difficulty})








def admin_2(request):
    role = request.session["role"]
    username = request.session['username']
    courses = Course.objects.all()
    processed_courses = []
    for course in courses:
        # print(course.title, course.description, course.instructor, course.image_url, course.video_url,course.detail_page_url,course.chapter)

        processed_course = {
            'title': course.title,
            'description': course.description,
            'instructor': course.instructor,
            'image_url': course.image_url,
            'video_url': course.video_url,
            'detail_page_url': course.detail_page_url,
            'chapter': course.chapter,
            'username': username,

        }
        print(processed_course)
        processed_courses.append(processed_course)

    return render(request, 'admin_2/index.html', {'courses': processed_courses, 'username': username})

def index_html(request):
    return render(request, 'admin_2/index.html')

def flex_layout_html(request):
    # 查询所有课程
    courses = Course.objects.all()

    # 提取所有学生信息
    all_students = []
    for course in courses:
        for i in range(1, 17):  # 遍历16个视频的学生信息
            students_field = getattr(course, f'video_{i}_students')
            if students_field:
                all_students.extend(students_field)

    # 渲染页面并将学生信息列表传递给模板
    return render(request, 'admin_2/flex-layout.html', {
        'students': all_students,
    })


def flow_layout_html(request):
    return render(request, 'admin_2/flow-layout.html')


def button_view(request):
    return render(request, 'admin_2/button.html')


def table_html(request):
    topics = Topic.objects.filter(module='teacher_qa')

    # 渲染页面并将topics列表传递给模板
    return render(request, 'admin_2/table.html', {
        'topics': topics,
    })


def form_html(request):
    return render(request, 'admin_2/form.html')


def popups_html(request):
    return render(request, 'admin_2/popups.html')

def echarts_html(request):
    return render(request, 'admin_2/echarts.html')


def echarts_html(request):
    return render(request, 'admin_2/echarts.html')
def ueditor_html(request):
    return render(request, 'admin_2/ueditor.html')

def progress_html(request):
    return render(request, 'admin_2/progress.html')


def tab_html(request):
    worst_students = Record.objects.order_by('grade')[:10]
    return render(request, 'admin_2/tab.html', {'worst_students': worst_students})


def button_dropdown_html(request):
    return render(request, 'admin_2/button-dropdown.html')


def title_html(request):
    return render(request, 'admin_2/title.html')


def paging_html(request):
    return render(request, 'admin_2/paging.html')


def animation_html(request):
    return render(request, 'admin_2/animation.html')

def breadcrumb_html(request):
    return render(request, 'admin_2/breadcrumb.html')



def paging_html(request):
    return render(request, 'admin_2/paging.html')


def download_and_grade(request, paper_id):
    # 获取对应的试卷
    paper = get_object_or_404(StudentPaper, id=paper_id)
    file_url = paper.file_url

    # 下载文件
    response = requests.get(file_url)
    temp_file_path = os.path.join(settings.MEDIA_ROOT, 'temp.docx')
    with open(temp_file_path, 'wb') as temp_file:
        temp_file.write(response.content)

    # 读取和批改试卷
    doc = docx.Document(temp_file_path)
    calculation_questions_answers = []

    question_count = 0
    score = 0
    correct_answers = "AABBCCDDAA"
    choice_question_index = 0

    for i, para in enumerate(doc.paragraphs):
        # 提取计算题答案和题干
        if '答：' in para.text and question_count < 5:
            answer_start = para.text.find('答：') + 2
            answer_end = para.text.find('（', answer_start) if '（' in para.text[answer_start:] else len(para.text)
            answer = para.text[answer_start:answer_end].strip()
            # 提取计算题的题干（假设题干在答案前一段）
            if i > 0:  # 确保前一个段落存在
                question = doc.paragraphs[i - 1].text.strip()
                calculation_questions_answers.append({'question': question, 'answer': answer})
            question_count += 1

        # 标记选择题答案
        if '（' in para.text and '）' in para.text and choice_question_index < 10:
            user_answer = para.text[para.text.find('（') + 1:para.text.find('）')]
            run = para.add_run(" ")
            if user_answer == correct_answers[choice_question_index]:
                run.text = "√"
                score += 2  # 每题2分
            else:
                run.text = "×"
            run.font.size = Pt(18)  # 设置字体大小
            run.font.color.rgb = RGBColor(255, 0, 0)  # 设置字体颜色为红色
            choice_question_index += 1

        # 标记填空题答案
        elif '______' in para.text:
            # 假设正确答案和学生答案已经存在变量中
            correct_fill_in_answer = "correct_answer"  # 示例答案
            student_answer = "student_answer"  # 示例学生答案

            run = para.add_run(" ")
            if student_answer == correct_fill_in_answer:
                run.text = "√"
                score += 2  # 填空题每题2分
            else:
                run.text = "×"
            run.font.size = Pt(30)
            run.font.color.rgb = RGBColor(255, 0, 0)

    # AI评分部分
    messages = [ChatMessage(
        role="user",
        content='假设你现在是一台阅卷机器，下面是保存着题目和回答的字典:' + str(calculation_questions_answers) + '请你根据题目，对答案评分1-10，只回答5道题的得分即5个数字即可，如果不够5道题，则其余的为0分。你的回答不能参杂文字，只能是5个数字，比如“4 5 7 3 0”'
    )]
    handler = ChunkPrintHandler()
    a = spark.generate([messages], callbacks=[handler])
    first_reply_content = a.generations[0][0].message.content
    print('AI的评分：', first_reply_content)

    # 将AI评分转换为列表，并计算总分
    ai_scores = list(map(int, first_reply_content.split()))
    ai_total_score = sum(ai_scores)
    score += ai_total_score  # 将AI评分总分加到总分中

    # 在简答题的空白处写入AI评分
    for i, para in enumerate(doc.paragraphs):
        if '答：' in para.text:
            para.add_run(f"\nAI评分总分: {ai_total_score}分").font.color.rgb = RGBColor(255, 0, 0)
            break

    # 在试卷右上方写入随机分数
    doc.paragraphs[0].add_run(f"\n\n批改得分: {score}分").bold = True

    # 保存修改后的试卷
    graded_file_path = os.path.join(settings.MEDIA_ROOT, f"graded_{os.path.basename(temp_file_path)}")
    doc.save(graded_file_path)

    # 打印提取的题干和答案（用于调试）
    for qa in calculation_questions_answers:
        print(f"题干: {qa['question']}")
        print(f"答案: {qa['answer']}")
    print('ssss', calculation_questions_answers)

    # 下载修改后的试卷
    with open(graded_file_path, 'rb') as f:
        response = HttpResponse(f.read(),
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={os.path.basename(graded_file_path)}'
        return response

    # 返回AI评分总分（用于其他用途）
    return ai_total_score



# def download_and_grade(request, paper_id):
#     # 获取对应的试卷
#     paper = StudentPaper.objects.get(id=paper_id)
#     file_url = paper.file_url
#
#
#     response = requests.get(file_url)
#     temp_file_path = os.path.join(settings.MEDIA_ROOT, 'temp.docx')
#     with open(temp_file_path, 'wb') as temp_file:
#         temp_file.write(response.content)
#
#     # 读取和批改试卷
#     doc = docx.Document(temp_file_path)
#     score = random.randint(60, 100)
#     bracket_count = 0
#     in_calculation_section = False
#
#     for para in doc.paragraphs:
#         text = para.text.strip()
#         # 选择题批改
#         if '（' in text and '）' in text:
#             bracket_count += 1
#             if bracket_count > 1:  # 从第二个括号开始批改
#                 if random.choice([True, False]):
#                     para.text += " √"
#                 else:
#                     para.text += " ×"
#         # 填空题批改
#         elif '______' in text:
#             if random.choice([True, False]):
#                 para.text += " √"
#             else:
#                 para.text += " ×"
#         # 检查是否进入计算题部分
#         elif '计算题' in text:
#             in_calculation_section = True
#         # 计算题批改
#         elif in_calculation_section and text:
#             if random.choice([True, False]):
#                 para.text += " √"
#             else:
#                 para.text += " ×"
#
#     # 在试卷右上方写入随机分数
#     doc.paragraphs[0].add_run(f"\n\n批改得分: {score}分").bold = True
#
#     # 保存修改后的试卷
#     graded_file_path = os.path.join(settings.MEDIA_ROOT, f"graded_{os.path.basename(temp_file_path)}")
#     doc.save(graded_file_path)
#
#     # 返回分数和下载链接
#     download_url = os.path.join(settings.MEDIA_URL, f"graded_{os.path.basename(temp_file_path)}")
#     return JsonResponse({'score': score, 'download_url': download_url})

def layer_html(request):
    papers = StudentPaper.objects.all()
    # paper_id = request.GET.get('paper_id')
    ai_total_score = download_and_grade(request, 1)
    return render(request, 'admin_2/layer.html', {'papers': papers, 'ai_total_score': ai_total_score})
    # return render(request, 'admin_2/layer.html', {'papers': papers})


def upexam(request):
    return render(request, 'upexam.html')



class Record:
    def __init__(self, name, grade):
        self.name = name
        self.grade = grade

# 后端函数，提供最差学生数据



#
#
# def tab_html(request):
#     worst_students = [
#         Record(name='张晓伟', grade=85),
#         Record(name='李娜', grade=80),
#         Record(name='王时磊', grade=78),
#         Record(name='赵敏', grade=76),
#         Record(name='孙洋', grade=75),
#         Record(name='周晓婷', grade=73),
#         Record(name='吴刚', grade=70),
#         Record(name='郑丽', grade=65),
#         Record(name='刘笑峰', grade=60),
#         Record(name='陈霞', grade=55),
#
#     ]
#     return render(request, 'admin_2/tab.html', {'worst_students': worst_students})


def get_warning_message(grade):
    if grade < 60:
        return '<span class="warning-retreat">退步</span>'
    elif 60 <= grade < 70:
        return '<span class="warning-decline">成绩下滑</span>'
    elif 70 <= grade < 80:
        return '<span class="warning-fluctuation">波动较大</span>'
    elif 80 <= grade < 90:
        return '<span class="warning-fluctuation">临界分数警告</span>'
    elif 90 <= grade <= 100:
        return '<span class="warning-cheating">疑似作弊</span>'
    else:
        return ''


def tab_html(request):
    worst_students = [
        {"name": "张晓伟", "grade": 85},
        {"name": "李娜", "grade": 80},
        {"name": "王时磊", "grade": 78},
        {"name": "赵敏", "grade": 76},
        {"name": "孙洋", "grade": 75},
        {"name": "周晓婷", "grade": 73},
        {"name": "吴刚", "grade": 70},
        {"name": "郑丽", "grade": 65},
        {"name": "刘笑峰", "grade": 60},
        {"name": "陈霞", "grade": 55},
    ]

    for student in worst_students:
        student['warning_message'] = get_warning_message(student['grade'])

    return render(request, 'admin_2/tab.html', {'worst_students': worst_students})















# 后端函数，提供校园暴力数据
# def get_violence_data(request):
#     violence_data = [
#         {
#             "lat": 31.2304,
#             "lon": 121.4737,
#             "location": "教学楼",
#             "student_names": ["张三", "李四"]
#         },
#         {
#             "lat": 31.2324,
#             "lon": 121.4757,
#             "location": "操场",
#             "student_names": ["王五", "赵六"]
#         }
#     ]
#     return JsonResponse(violence_data, safe=False)



# def get_violence_data(request):
#     comments = Comment.objects.all()  # 获取所有评论，假设都与校园暴力相关
#     violence_data = []
#     for comment in comments:
#         violence_data.append({
#             "lat": 30.5460,  # 固定位置，实际可根据需要动态生成
#             "lon": 114.3000,  # 固定位置，实际可根据需要动态生成
#             "location": "教学楼",  # 固定位置，实际可根据需要动态生成
#             "student_name": comment.student_name,
#             "content": comment.content
#         })
#     print(violence_data)
#     return JsonResponse(violence_data, safe=False)



def get_violence_data(request):
    comments = Comment.objects.all()  # 获取所有评论，假设都与校园暴力相关
    base_lat = 31.941567
    base_lon = 118.788469
    violence_data = []
    for comment in comments:
        messages = [ChatMessage(
            role="user",
            content=  comment.content
        )]
        handler = ChunkPrintHandler()
        a = spark.generate([messages], callbacks=[handler])
        first_reply_content = a.generations[0][0].message.content
        print('AI鉴定结果', first_reply_content)
        # 随机生成一个位置，模拟每个评论在不同位置


        lat_offset = random.uniform(-0.01, 0.01)
        lon_offset = random.uniform(-0.01, 0.01)
        if first_reply_content != "[NO]":
            try:
                # 转换中文逗号为英文逗号
                first_reply_content = first_reply_content.replace('，', ',')
                # 解析鉴定结果
                location, student_name, people_involved = first_reply_content.strip("[]").split(",")

                # 随机生成纬度和经度偏移，确保定位点随机
                lat_offset = random.uniform(-0.0005, 0.0005)
                lon_offset = random.uniform(-0.0005, 0.0005)

                violence_data.append({
                    "lat": 31.941567 + lat_offset,
                    "lon": 	118.788469 + lon_offset,
                    "location": location,
                    "student_name": student_name,
                    "people_involved": people_involved,
                    "content": comment.content
                })
            except ValueError:
                # 处理解析错误的情况
                print(f"Invalid AI result format: {first_reply_content}")

    return JsonResponse(violence_data, safe=False)




logger = logging.getLogger(__name__)

@csrf_exempt
def send_alert(request):
    if request.method == 'POST':
        alert_type = request.POST.get('type')
        student = request.POST.get('student')
        content = request.POST.get('content')
        location = request.POST.get('location')

        # 打印日志以调试
        logger.info(f"Alert Type: {alert_type}")
        logger.info(f"Student: {student}")
        logger.info(f"Content: {content}")
        logger.info(f"Location: {location}")

        # 构建邮件内容
        subject = f"预警短信: 可能存在校园暴力"
        message = f"学生: {student}\n地点: {location}\n该学生发布内容: {content}\n可能存在校园暴力行为，请及时查看！"

        # 接收者的邮箱地址
        recipient_list = ['2451909851@qq.com']  # 替换为实际的接收者邮箱地址

        try:
            # 发送邮件
            send_mail(subject, message, settings.EMAIL_HOST_USER, recipient_list)
            return JsonResponse({'message': f'已向{alert_type}发送预警邮件'}, status=200)
        except Exception as e:
            logger.error(f"Error sending email: {e}")
            return JsonResponse({'error': '邮件发送失败'}, status=500)

    return JsonResponse({'error': 'Invalid request'}, status=400)






#
#
# ml_curriculum = Curriculum(name="机器学习", description="机器学习基础课程").save()
#
# # 添加知识点
# linear_regression = KnowledgePoint(name="线性回归", description="线性回归是一种监督学习方法...", grade_level=1).save()
#
# # 添加子知识点
# gradient_descent = KnowledgePoint(name="梯度下降", description="梯度下降是一种优化算法...", grade_level=2).save()
#
# # 将知识点与课程关联
# ml_curriculum.curriculum.connect(linear_regression)
#
# # 创建知识点之间的关系
# linear_regression.parent.connect(gradient_descent)








graph = Graph("bolt://localhost:7687", auth=("neo4j", "666666"))

# def add_knowledge_point(name, description, grade_level):
#     knowledge_point = Node("KnowledgePoint", name=name, description=description, grade_level=grade_level)
#     graph.create(knowledge_point)
#     print(f"知识点{name}添加成功")




# def add_course(name, description):
#     course = Node("Course", name=name, description=description)
#     graph.create(course)
#     return course
#
# # 创建知识点节点并与课程关联
# def add_knowledge_point(course, name, description, grade_level):
#     knowledge_point = Node("KnowledgePoint", name=name, description=description, grade_level=grade_level)
#     graph.create(knowledge_point)
#     rel = Relationship(course, "CONTAINS", knowledge_point)
#     graph.create(rel)
#     return knowledge_point
#
# # 创建知识点之间的关系
# def relate_knowledge_points(kp1, kp2, relationship_type):
#     rel = Relationship(kp1, relationship_type, kp2)
#     graph.create(rel)
#
#
# def get_all_knowledge_points():
#     query = "MATCH (kp:KnowledgePoint) RETURN kp"
#     results = graph.run(query)
#     for record in results:
#         print(record["kp"])


def add_course(name, description):
    course = Node("Course", name=name, description=description)
    graph.merge(course, "Course", "name")  # 使用 merge 确保不重复创建课程节点
    return course

# 创建知识点节点并与课程关联
def add_knowledge_point(course, name, description, grade_level):
    knowledge_point = Node("KnowledgePoint", name=name, description=description, grade_level=grade_level)
    graph.merge(knowledge_point, "KnowledgePoint", "name")  # 使用 merge 确保不重复创建知识点节点
    rel = Relationship(course, "CONTAINS", knowledge_point)
    graph.merge(rel)  # 确保关系也不会重复
    return knowledge_point

# 创建知识点之间的关系
def relate_knowledge_points(kp1, kp2, relationship_type):
    rel = Relationship(kp1, relationship_type, kp2)
    graph.merge(rel)  # 使用 merge 确保关系不重复

def get_all_knowledge_points():
    query = "MATCH (kp:KnowledgePoint) RETURN kp"
    results = graph.run(query)
    for record in results:
        print(record["kp"])





def exam(request,course_id,chapter,announcement_id):
    examid = course_id * 11
    course = Course.objects.get(id=course_id)
    processed_course = {
        'title': course.title,
        'description': course.description,
        'instructor': course.instructor,
        'image_url': course.image_url,
        'video_url': course.video_url,
        'detail_page_url': course.detail_page_url,
        'chapter': course.chapter,
    }
    # add_knowledge_point("机器学习", "线性回归是一种监督学习方法...", 1)

    # ai_course = add_course("人工智能课程", "这门课程涵盖了人工智能的基础概念和高级主题。")
    # nn_kp = add_knowledge_point(ai_course, "神经网络", "介绍神经网络的基本概念和应用。", 2)
    # perceptron_kp = add_knowledge_point(ai_course, "感知器模型", "感知器是一种最简单的神经网络模型。", 3)
    # activation_kp = add_knowledge_point(ai_course, "激活函数", "激活函数用于引入非线性特性。", 3)
    # # 建立知识点之间的关系
    # relate_knowledge_points(nn_kp, perceptron_kp, "CONTAINS")
    # relate_knowledge_points(nn_kp, activation_kp, "CONTAINS")
    # print("查看知识点")
    # get_all_knowledge_points()









    template_name = f'exam_{examid}.html'
    return render(request, template_name, {'course': processed_course})







def loudongjiance(request):

    role = request.session["role"]
    username = request.session['username']
    courses = Course.objects.all()
    processed_courses = []
    for course in courses:
        # print(course.title, course.description, course.instructor, course.image_url, course.video_url,course.detail_page_url,course.chapter)

        processed_course = {
            'title': course.title,
            'description': course.description,
            'instructor': course.instructor,
            'image_url': course.image_url,
            'video_url': course.video_url,
            'detail_page_url': course.detail_page_url,
            'chapter': course.chapter,
            'username': username,

        }
        print(processed_course)
        processed_courses.append(processed_course)

    # return render(request, 'ceshi_2/index.html', {'courses': processed_courses, 'username': username})

    return render(request, 'neo4j.html', {'courses': processed_courses, 'username': username})

graph = Graph('bolt://localhost:7687', auth=('neo4j', '666666'))

def search_all():
    data = []
    links = []

    # 查询所有节点
    nodes = graph.run("MATCH (n) RETURN n").data()
    for record in nodes:
        node = record['n']
        dict_node = {
            'name': node['name'],  # 确保节点名称是字符串
            'symbolSize': 50,
            'category': '对象'
        }
        data.append(dict_node)

    # 查询所有关系
    relationships = graph.run("MATCH (n)-[r]->(m) RETURN n, r, m").data()
    for record in relationships:
        source = record['n']['name']  # 确保名称是字符串
        target = record['m']['name']  # 确保名称是字符串
        name = record['r'].__class__.__name__  # 获取关系类型名称

        dict_link = {
            'source': source,
            'target': target,
            'name': name
        }
        links.append(dict_link)

    neo4j_data = {
        'data': data,
        'links': links
    }

    return neo4j_data




def neo4j_visualization(request):
    neo4j_data = search_all()

    print('后端的neo4j数据', neo4j_data)  # 检查后端返回的数据
    return render(request, 'neo4j.html', {
        'nodes_data': json.dumps(neo4j_data['data'], ensure_ascii=False),
        'links_data': json.dumps(neo4j_data['links'], ensure_ascii=False)
    })




def get_graph_data(request):
    data = []
    links = []

    # 查询所有节点
    nodes = graph.run("MATCH (n) RETURN n").data()
    for record in nodes:
        node = record['n']
        dict_node = {
            'id': node['name'],  # 节点 ID
            'name': node['name'],  # 节点名称
            'label': '对象',  # 节点类别
        }
        data.append(dict_node)

    # 查询所有关系
    relationships = graph.run("MATCH (n)-[r]->(m) RETURN n, r, m").data()
    for record in relationships:
        dict_link = {
            'source': record['n']['name'],  # 起始节点 ID
            'target': record['m']['name'],  # 目标节点 ID
            'label': record['r'].__class__.__name__  # 关系类型
        }
        links.append(dict_link)

    return JsonResponse({'nodes': data, 'links': links})


def graph_visualization(request):
    questions = QuizQuestion.objects.all()
    total_questions = questions.count()


    # ================================================================================================
    username = '周'
    if username:
        student_topics = Topic.objects.filter(student_name=username)

        # 统计每个课程中的话题数量
        topic_counts = {}
        for topic in student_topics:
            course_title = topic.course.title
            if course_title in topic_counts:
                topic_counts[course_title] += 1
            else:
                topic_counts[course_title] = 1

    if not username:
        return JsonResponse({'status': 'fail', 'error': 'user not logged in'}, status=400)

    course_id = request.GET.get('course_id', 1)  # 默认显示课程ID为1
    student_data = get_student_data(username, course_id)
    all_courses = Course.objects.all()  # 获取所有课程列表

    # 初始化变量
    total_likes = 0
    total_favorites = 0
    total_score = 0
    comment_dict = {}  # 存储所有评论的字典
    sentiment_scores = {}  # 存储评论的情感得分
    progress_list = [0] * 16  # 存储16个视频的播放进度
    useful_progress_number = 0  # 播放进度大于30的个数
    test_scores = [0] * 5  # 存储5门测试的具体分数
    test_participation_count = 0  # 参与考试的数量
    num_videos = 16
    num_tests = 5

    # 计算点赞量、收藏量和评论
    for video in student_data[0]['videos']:
        video_id = video.get('video_id', 0)  # 获取video_id
        liked = video.get('liked', 0)
        favorited = video.get('favorited', 0)
        progress = video.get('progress', 0)
        comment = video.get('comment', '评论未录入')

        # 计算总点赞量和收藏量
        if liked == 1:
            total_likes += 1
        if favorited == 1:
            total_favorites += 1

        # 更新视频播放进度列表
        if 0 <= video_id - 1 < num_videos:  # Assuming video_id starts from 1
            progress_list[video_id - 1] = progress

        # 计算进度大于30的个数
        if isinstance(progress, int) and progress > 30:
            useful_progress_number += 1

        # 收集评论，将评论加入字典，使用video_id作为键
        comment_dict[video_id] = comment
    print('pl', comment_dict)

    messages = [ChatMessage(
        role="user",
        content= str(comment_dict)
    )]
    handler = ChunkPrintHandler()
    a = spark.generate([messages], callbacks=[handler])
    first_reply_content = a.generations[0][0].message.content
    print('AI鉴定偏好结果', first_reply_content)
    print('leix a ', type(first_reply_content))
    # 根据AI返回的偏好结果生成 preference_list
    # 处理 AI 返回的偏好结果，将其分割成列表
    preference_list = [item.strip() for item in first_reply_content.split('，') if item.strip()]

    # 处理 preference_list，将其转换成字典

    # 确保 preference_list 只包含前三个偏好
    preference_list = preference_list[:3]
    print('偏好列表', preference_list)
    uncorrect_knowledge_points = ['FTP', '被动模式', '主动模式']

    video_urls = [
        "https://analysis-zhouyujie.oss-cn-hangzhou.aliyuncs.com/pianhaovideo/1.mp4",
        "https://analysis-zhouyujie.oss-cn-hangzhou.aliyuncs.com/pianhaovideo/2.mp4",
        "https://analysis-zhouyujie.oss-cn-hangzhou.aliyuncs.com/pianhaovideo/4.mp4",
    ]
# ===================================================================================================================
    video_preferences = list(zip(video_urls, preference_list))



    context = {
        'questions': questions,
        'total_questions': total_questions,
        'video_preferences': video_preferences,
        'uncorrect_knowledge_points': uncorrect_knowledge_points,  # 添加这行代码

    }

    return render(request, 'ceshi_2/index.html', context)


def check_answers(request):
    if request.method == 'POST':
        # 处理提交的答案
        submitted_answers = request.POST.getlist('answers')  # 假设前端传递了答案列表
        print('答案列表',submitted_answers)
        submitted_answers=['A','B','C','A','C']

        # 初始化一个错误的知识点列表
        wrong_knowledge_points = []

        for answer_id, submitted_answer in enumerate(submitted_answers):
            try:
                question = QuizQuestion.objects.get(id=answer_id + 1)
                if question.correct_answer != submitted_answer:
                    # 如果回答错误，将相关的知识点添加到列表中
                    wrong_knowledge_points.append(question.knowledge_point.name)
            except QuizQuestion.DoesNotExist:
                return JsonResponse({"error": "Invalid question ID"}, status=400)
        print('check answers函数错误的知识点', wrong_knowledge_points)


        # 返回包含错误知识点的 JSON 响应
        return JsonResponse({
            "wrong_knowledge_points": wrong_knowledge_points
        })

    # 如果不是 POST 请求，返回一个错误响应
    return JsonResponse({"error": "Invalid request method"}, status=405)




def save_elements(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            # 处理并保存 data 中的元素
            # 例如，保存到数据库，或者在日志中记录
            # print("错误知识点列表:", data)
            unique_labels = set()

            for item in data:
                label = item['data']['label']
                if '-' not in label and not label[0].isdigit() and not label.startswith(('第', '三','知','选')):
                    unique_labels.add(label)

            # 转换为列表并输出
            result = list(unique_labels)
            print('漏洞知识点',result)
            # 可以在这里添加代码将数据保存到数据库中
            # ...

            return JsonResponse({"message": "Elements saved successfully"})
        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON"}, status=400)
    else:
        return JsonResponse({"error": "Invalid request method"}, status=405)





def submit_student_answer(request):
    if request.method == 'POST':
        topic_id = request.POST.get('topic_id')
        answer_content = request.POST.get('answer')

        # 获取对应的问题（Topic）
        topic = get_object_or_404(Topic, id=topic_id)

        # 获取当前用户的 Student 对象


        # 硬编码的数据部分
        fixed_author_id = '1'  # 假设为当前用户的学生ID
        fixed_topic_id = 74  # 预设的 topic_id
        fixed_parent_comment_id = 98  # 预设的 parent_comment_id，可以为 None
        fixed_created_at = datetime(2024, 8, 18, 14, 30)  # 预设的创建时间
        fixed_updated_at = fixed_created_at  # 通常与创建时间一致
        fixed_student_name = '周涛'  # 假设学生有 name 属性

        # 创建并保存 Comment 实例
        comment = Comment.objects.create(
            content=answer_content,  # 用户提供的回答内容
            author_id=fixed_author_id,
            topic_id=fixed_topic_id,
            parent_comment_id=fixed_parent_comment_id,
            created_at=fixed_created_at,
            updated_at=fixed_updated_at,
            student_name=fixed_student_name
        )

        # 返回成功的 JSON 响应
        return JsonResponse({'status': 'success', 'message': '回答成功！'})

        # 如果请求方法不是 POST，则返回错误信息
    return JsonResponse({'status': 'error', 'message': '无效的请求'}, status=400)





def AI_reply_student_question(request):
    if request.method == 'POST':
        topic_id = request.POST.get('topic_id')

        # 获取对应的问题（Topic）
        topic = get_object_or_404(Topic, id=topic_id)

        # 在这里调用你的AI回答逻辑，模拟返回AI的回答内容

        messages = [ChatMessage(
            role="user",
            content='先清空之前的记忆，回答下面学生提出来的问题,直接输出答案即可，50字以内'+topic.title
        )]
        handler = ChunkPrintHandler()
        a = spark.generate([messages], callbacks=[handler])
        first_reply_content = a.generations[0][0].message.content
        print('AI对学生问题回答的结果：', first_reply_content)



        # 返回成功的 JSON 响应
        return JsonResponse({'status': 'success', 'first_reply_content': first_reply_content})

    # 如果请求方法不是 POST，则返回错误信息
    return JsonResponse({'status': 'error', 'message': '无效的请求'}, status=400)














